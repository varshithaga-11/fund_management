import os
import random
import logging
from datetime import datetime
from decimal import Decimal

from django.conf import settings
from django.contrib.auth.hashers import make_password
from django.db import transaction
from django.http import FileResponse, HttpResponse
from django.shortcuts import render
from django.utils import timezone

from rest_framework import viewsets, status
from rest_framework.generics import RetrieveUpdateDestroyAPIView
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.response import Response
from rest_framework.views import APIView

from openpyxl import load_workbook, Workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber
from io import BytesIO

from .models import *
from .serializers import *

logger = logging.getLogger(__name__)



# Create your views here.
class UserRegisterView(APIView):
    permission_classes = [AllowAny]  
    
    def post(self, request):
        try:
            serializer = UserRegisterSerializer(data=request.data)
            if serializer.is_valid():
                user = serializer.save()
                return Response({
                    "status": "success",
                    "response_code": status.HTTP_201_CREATED,
                    "message": "User registered successfully",
                    "user": {
                        "id": user.id,
                        "username": user.username,
                        "email": user.email,
                        "role": user.role,
                        "created_by": user.created_by.id if user.created_by else None
                    }
                })
            return Response({
                "status": "failed",
                "response_code": status.HTTP_400_BAD_REQUEST,
                "message": serializer.errors
            })
        except Exception as e:
            message = str(e)
            return Response({
                "status": "failed",
                "response_code": status.HTTP_500_INTERNAL_SERVER_ERROR,
                "message": message
            })
 

class LoginView(APIView):
    permission_classes = [AllowAny] 
    
    def post(self, request):
        try:
            serializer = LoginSerializer(data=request.data)
            if serializer.is_valid():
                return Response({"status":"success","response_code":status.HTTP_200_OK,"message":"User logged in successfully","tokens":serializer.validated_data})
            return Response({"status":"failed","response_code":status.HTTP_404_NOT_FOUND,"message":serializer.errors})
        except Exception as e:
            message = str(e)
            return Response({"status":"failed","response_code":status.HTTP_500_INTERNAL_SERVER_ERROR,"message":message})
        

class RefreshTokenView(APIView):
    permission_classes = [AllowAny] 

    def post(self, request):
        try:
            serializer = RefreshTokenSerializer(data=request.data)
            if serializer.is_valid():
                return Response(serializer.validated_data)
            return Response({"status":"failed","response_code":status.HTTP_404_NOT_FOUND,"message":serializer.errors})
        except Exception as e:
            message = str(e)
            return Response({"status":"failed","response_code":status.HTTP_500_INTERNAL_SERVER_ERROR,"message":message})


class ProfileView(viewsets.ModelViewSet):
    # permission_classes = [IsAuthenticated]
    queryset = UserRegister.objects.all()
    serializer_class = ProfileSerializer





class CompanyViewSet(viewsets.ModelViewSet):
    queryset = Company.objects.all().order_by("-created_at")
    serializer_class = CompanySerializer
    permission_classes = [IsAuthenticated]


class FinancialPeriodViewSet(viewsets.ModelViewSet):
    queryset = FinancialPeriod.objects.all().order_by("-created_at")
    serializer_class = FinancialPeriodSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = FinancialPeriod.objects.all()
        company_id = self.request.query_params.get('company', None)
        if company_id:
            queryset = queryset.filter(company_id=company_id)
        return queryset.order_by("-created_at")


class TradingAccountViewSet(viewsets.ModelViewSet):
    queryset = TradingAccount.objects.all()
    serializer_class = TradingAccountSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = TradingAccount.objects.all()
        period_id = self.request.query_params.get('period', None)
        if period_id:
            queryset = queryset.filter(period_id=period_id)
        return queryset


class ProfitAndLossViewSet(viewsets.ModelViewSet):
    queryset = ProfitAndLoss.objects.all()
    serializer_class = ProfitAndLossSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = ProfitAndLoss.objects.all()
        period_id = self.request.query_params.get('period', None)
        if period_id:
            queryset = queryset.filter(period_id=period_id)
        return queryset


class BalanceSheetViewSet(viewsets.ModelViewSet):
    queryset = BalanceSheet.objects.all()
    serializer_class = BalanceSheetSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = BalanceSheet.objects.all()
        period_id = self.request.query_params.get('period', None)
        if period_id:
            queryset = queryset.filter(period_id=period_id)
        return queryset


class OperationalMetricsViewSet(viewsets.ModelViewSet):
    queryset = OperationalMetrics.objects.all()
    serializer_class = OperationalMetricsSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = OperationalMetrics.objects.all()
        period_id = self.request.query_params.get('period', None)
        if period_id:
            queryset = queryset.filter(period_id=period_id)
        return queryset


class RatioResultViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = RatioResult.objects.all()
    serializer_class = RatioResultSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = RatioResult.objects.all()
        period_id = self.request.query_params.get('period', None)
        if period_id:
            queryset = queryset.filter(period_id=period_id)
        return queryset


class StatementColumnConfigViewSet(viewsets.ModelViewSet):
    """
    Manage display names / order for financial statement columns.
    Supports global configs (company null) and company-specific overrides.
    """
    queryset = StatementColumnConfig.objects.all()
    serializer_class = StatementColumnConfigSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = StatementColumnConfig.objects.all()
        company_id = self.request.query_params.get("company")
        statement_type = self.request.query_params.get("statement_type")
        if company_id == "global":
            qs = qs.filter(company__isnull=True)
        elif company_id:
            qs = qs.filter(company_id=company_id)
        # if no company param: return all (caller can filter client-side)
        if statement_type:
            qs = qs.filter(statement_type=statement_type)
        return qs.order_by("canonical_field")


class CalculateRatiosView(APIView):
    permission_classes = [IsAuthenticated]
    
    def post(self, request, period_id=None):
        """
        Calculate ratios for a given period
        POST /api/periods/<period_id>/calculate-ratios/
        """
        try:
            if period_id is None:
                period_id = request.data.get('period_id')
            
            if not period_id:
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "period_id is required"
                })
            
            period = FinancialPeriod.objects.get(id=period_id)
            
            # Validate all required data exists
            if not hasattr(period, 'trading_account'):
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "TradingAccount not found for this period"
                })
            if not hasattr(period, 'profit_loss'):
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "ProfitAndLoss not found for this period"
                })
            if not hasattr(period, 'balance_sheet'):
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "BalanceSheet not found for this period"
                })
            if not hasattr(period, 'operational_metrics'):
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "OperationalMetrics not found for this period"
                })
            
            # Calculate ratios
            from app.services.ratio_calculator import RatioCalculator
            from decimal import Decimal
            
            calculator = RatioCalculator(period)
            all_ratios = calculator.calculate_all_ratios()
            traffic_light_statuses = calculator.get_traffic_light_statuses()
            
            # Create or update RatioResult
            ratio_result, created = RatioResult.objects.get_or_create(
                period=period,
                defaults={
                    'working_fund': Decimal(str(all_ratios['working_fund'])),
                    'stock_turnover': Decimal(str(all_ratios.get('stock_turnover', 0))),
                    'gross_profit_ratio': Decimal(str(all_ratios.get('gross_profit_ratio', 0))),
                    'net_profit_ratio': Decimal(str(all_ratios.get('net_profit_ratio', 0))),
                    'own_fund_to_wf': Decimal(str(all_ratios.get('own_fund_to_wf', 0))),
                    'deposits_to_wf': Decimal(str(all_ratios.get('deposits_to_wf', 0))),
                    'borrowings_to_wf': Decimal(str(all_ratios.get('borrowings_to_wf', 0))),
                    'loans_to_wf': Decimal(str(all_ratios.get('loans_to_wf', 0))),
                    'investments_to_wf': Decimal(str(all_ratios.get('investments_to_wf', 0))),
                    'cost_of_deposits': Decimal(str(all_ratios.get('cost_of_deposits', 0))),
                    'yield_on_loans': Decimal(str(all_ratios.get('yield_on_loans', 0))),
                    'yield_on_investments': Decimal(str(all_ratios.get('yield_on_investments', 0))),
                    'credit_deposit_ratio': Decimal(str(all_ratios.get('credit_deposit_ratio', 0))),
                    'avg_cost_of_wf': Decimal(str(all_ratios.get('avg_cost_of_wf', 0))),
                    'avg_yield_on_wf': Decimal(str(all_ratios.get('avg_yield_on_wf', 0))),
                    'gross_fin_margin': Decimal(str(all_ratios.get('gross_fin_margin', 0))),
                    'operating_cost_to_wf': Decimal(str(all_ratios.get('operating_cost_to_wf', 0))),
                    'net_fin_margin': Decimal(str(all_ratios.get('net_fin_margin', 0))),
                    'risk_cost_to_wf': Decimal(str(all_ratios.get('risk_cost_to_wf', 0))),
                    'net_margin': Decimal(str(all_ratios.get('net_margin', 0))),
                    'all_ratios': all_ratios,
                    'traffic_light_status': traffic_light_statuses
                }
            )
            
            if not created:
                # Update existing
                ratio_result.working_fund = Decimal(str(all_ratios['working_fund']))
                ratio_result.stock_turnover = Decimal(str(all_ratios.get('stock_turnover', 0)))
                ratio_result.gross_profit_ratio = Decimal(str(all_ratios.get('gross_profit_ratio', 0)))
                ratio_result.net_profit_ratio = Decimal(str(all_ratios.get('net_profit_ratio', 0)))
                ratio_result.own_fund_to_wf = Decimal(str(all_ratios.get('own_fund_to_wf', 0)))
                ratio_result.deposits_to_wf = Decimal(str(all_ratios.get('deposits_to_wf', 0)))
                ratio_result.borrowings_to_wf = Decimal(str(all_ratios.get('borrowings_to_wf', 0)))
                ratio_result.loans_to_wf = Decimal(str(all_ratios.get('loans_to_wf', 0)))
                ratio_result.investments_to_wf = Decimal(str(all_ratios.get('investments_to_wf', 0)))
                ratio_result.cost_of_deposits = Decimal(str(all_ratios.get('cost_of_deposits', 0)))
                ratio_result.yield_on_loans = Decimal(str(all_ratios.get('yield_on_loans', 0)))
                ratio_result.yield_on_investments = Decimal(str(all_ratios.get('yield_on_investments', 0)))
                ratio_result.credit_deposit_ratio = Decimal(str(all_ratios.get('credit_deposit_ratio', 0)))
                ratio_result.avg_cost_of_wf = Decimal(str(all_ratios.get('avg_cost_of_wf', 0)))
                ratio_result.avg_yield_on_wf = Decimal(str(all_ratios.get('avg_yield_on_wf', 0)))
                ratio_result.gross_fin_margin = Decimal(str(all_ratios.get('gross_fin_margin', 0)))
                ratio_result.operating_cost_to_wf = Decimal(str(all_ratios.get('operating_cost_to_wf', 0)))
                ratio_result.net_fin_margin = Decimal(str(all_ratios.get('net_fin_margin', 0)))
                ratio_result.risk_cost_to_wf = Decimal(str(all_ratios.get('risk_cost_to_wf', 0)))
                ratio_result.net_margin = Decimal(str(all_ratios.get('net_margin', 0)))
                ratio_result.all_ratios = all_ratios
                ratio_result.traffic_light_status = traffic_light_statuses
                ratio_result.save()
            
            serializer = RatioResultSerializer(ratio_result)
            return Response({
                "status": "success",
                "response_code": status.HTTP_200_OK,
                "message": "Ratios calculated successfully",
                "data": serializer.data
            })
            
        except FinancialPeriod.DoesNotExist:
            return Response({
                "status": "failed",
                "response_code": status.HTTP_404_NOT_FOUND,
                "message": "FinancialPeriod not found"
            })
        except Exception as e:
            return Response({
                "status": "failed",
                "response_code": status.HTTP_500_INTERNAL_SERVER_ERROR,
                "message": str(e)
            })


class UploadExcelView(APIView):
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        """
        Upload Excel file and parse all 4 sheets
        POST /api/upload-excel/
        """
        logger.info("=== UploadExcelView POST request received ===")
        try:
            if 'file' not in request.FILES:
                logger.warning("DEBUG: No file provided")
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "No file provided"
                })
            
            company_id = request.data.get('company_id')
            if not company_id:
                logger.warning("DEBUG: company_id is required")
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "company_id is required"
                })
            
            try:
                company = Company.objects.get(id=company_id)
                logger.info(f"DEBUG: Found company - {company.name}")
            except Company.DoesNotExist:
                logger.warning(f"DEBUG: Company not found - ID: {company_id}")
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_404_NOT_FOUND,
                    "message": "Company not found"
                })
            
            uploaded_file = request.FILES['file']
            logger.info(f"DEBUG: File received - {uploaded_file.name}")
            filename = uploaded_file.name
            period_info = self._extract_period_from_filename(filename)

            ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
            if ext not in ('xlsx', 'xls', 'docx', 'pdf'):
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "Unsupported file type. Use .xlsx, .xls, .docx, or .pdf"
                })

            # Handle .docx or .pdf: store file and create period with default empty data
            if ext == 'docx':
                period = self._create_period_from_document(
                    request, company, uploaded_file, period_info, file_type='docx'
                )
                return Response({
                    "status": "success",
                    "response_code": status.HTTP_200_OK,
                    "message": "Document (.docx) uploaded successfully",
                    "period_id": period.id,
                    "period_label": period.label
                })
            if ext == 'pdf':
                period = self._create_period_from_document(
                    request, company, uploaded_file, period_info, file_type='pdf'
                )
                return Response({
                    "status": "success",
                    "response_code": status.HTTP_200_OK,
                    "message": "Document (.pdf) uploaded successfully",
                    "period_id": period.id,
                    "period_label": period.label
                })

            # Excel path
            excel_file = uploaded_file
            # Load workbook
            logger.info("DEBUG: Loading workbook...")
            workbook = load_workbook(excel_file, data_only=True)
            logger.info(f"DEBUG: Workbook loaded, sheets: {workbook.sheetnames}")
            
            # Find required sheets – support both formats
            sheet_mapping = self._find_sheets(workbook)
            logger.info(f"DEBUG: Found sheets in mapping: {list(sheet_mapping.keys())}")
            available = workbook.sheetnames

            # Format A: Financial_Statement, Balance_Sheet_Liabilities, Balance_Sheet_Assets, Profit_Loss, Trading_Account
            format_a = all(k in sheet_mapping for k in ['Financial_Statement', 'Balance_Sheet_Liabilities', 'Balance_Sheet_Assets', 'Profit_Loss', 'Trading_Account'])

            # Format B: Balance Sheet, Profit and Loss, Trading Account, Operational Metrics
            format_b = all(k in sheet_mapping for k in ['Balance Sheet', 'Profit and Loss', 'Trading Account', 'Operational Metrics'])
            
            logger.info(f"DEBUG: Format A detected: {format_a}, Format B detected: {format_b}")

            if format_a:
                logger.info("DEBUG: Parsing Format A sheets...")
                financial_statement_data = self._parse_financial_statement_sheet(workbook[sheet_mapping['Financial_Statement']])
                liabilities_data = self._parse_balance_sheet_liabilities(workbook[sheet_mapping['Balance_Sheet_Liabilities']], company)
                assets_data = self._parse_balance_sheet_assets(workbook[sheet_mapping['Balance_Sheet_Assets']], company)
                balance_sheet_data = self._default_balance_sheet({**liabilities_data, **assets_data})
                profit_loss_data = self._default_profit_loss(self._parse_profit_loss_rows(workbook[sheet_mapping['Profit_Loss']], company))
                trading_account_data = self._default_trading_account(self._parse_trading_account_rows(workbook[sheet_mapping['Trading_Account']], company))
                staff_count = financial_statement_data.get('staff_count')
                if staff_count is not None:
                    operational_metrics_data = {'staff_count': int(float(staff_count))}
                else:
                    operational_metrics_data = {'staff_count': 1}
                fiscal_end = financial_statement_data.get('fiscal_year_end')
                if fiscal_end and not period_info.get('label'):
                    period_info['label'] = fiscal_end
                    period_info['end_date'] = fiscal_end
            elif format_b:
                logger.info("DEBUG: Parsing Format B sheets...")
                balance_sheet_data = self._default_balance_sheet(self._parse_balance_sheet(workbook[sheet_mapping['Balance Sheet']], company))
                profit_loss_data = self._default_profit_loss(self._parse_profit_loss(workbook[sheet_mapping['Profit and Loss']], company))
                trading_account_data = self._default_trading_account(self._parse_trading_account(workbook[sheet_mapping['Trading Account']], company))
                operational_metrics_data = self._parse_operational_metrics(workbook[sheet_mapping['Operational Metrics']], company)
                logger.info("DEBUG: Format B sheets parsed successfully")
            elif len(available) == 1 and available[0] == 'Sheet':
                # Format C: Single generic "Sheet" - try to auto-detect and parse as balance sheet
                logger.info("DEBUG: Detecting single generic 'Sheet'... attempting to auto-parse as balance sheet")
                single_sheet = workbook[available[0]]
                
                # Try to detect if it's balance sheet data by checking headers
                first_row = [str(cell.value).strip().lower() if cell.value else '' for cell in next(single_sheet.iter_rows(values_only=False))]
                has_liabilities = any('liabilit' in h for h in first_row)
                has_assets = any('asset' in h for h in first_row)
                
                if has_liabilities and has_assets:
                    logger.info("DEBUG: Detected balance sheet format - parsing as Format B (single sheet)")
                    balance_sheet_data = self._parse_balance_sheet(single_sheet, company)
                    # Set default/empty data for other required sheets
                    profit_loss_data = self._default_profit_loss({})
                    trading_account_data = self._default_trading_account({})
                    operational_metrics_data = {'staff_count': 1}
                else:
                    logger.error(f"DEBUG: Could not auto-detect sheet type. First row: {first_row}")
                    return Response({
                        "status": "failed",
                        "response_code": status.HTTP_400_BAD_REQUEST,
                        "message": f"Could not parse single sheet format. Please use properly named sheets or ensure balance sheet has 'Liabilities' and 'Assets' headers."
                    })
            else:
                logger.error(f"DEBUG: Unsupported sheet set. Available sheets: {', '.join(available)}")
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": f"Unsupported sheet set. Use either: (1) Financial_Statement, Balance_Sheet_Liabilities, Balance_Sheet_Assets, Profit_Loss, Trading_Account OR (2) Balance Sheet, Profit and Loss, Trading Account, Operational Metrics. Available: {', '.join(available)}"
                })

            # Create Financial Period (filename e.g. April_2025, or fiscal year from Financial_Statement)
            period_label = request.data.get('period_label') or period_info.get('label') or f"FY-{datetime.now().year}-{datetime.now().year + 1}"
            start_date = request.data.get('start_date') or period_info.get('start_date') or f"{datetime.now().year}-04-01"
            end_date = request.data.get('end_date') or period_info.get('end_date') or f"{datetime.now().year + 1}-03-31"
            period_type = request.data.get('period_type') or period_info.get('period_type') or ('MONTHLY' if period_info else 'YEARLY')
            
            logger.info(f"DEBUG: Creating/updating period - Company: {company.name}, Label: {period_label}")
            
            # Use transaction to ensure all data is saved together
            with transaction.atomic():
                period, created = FinancialPeriod.objects.get_or_create(
                    company=company,
                    label=period_label,
                    defaults={
                        'period_type': period_type,
                        'start_date': start_date,
                        'end_date': end_date,
                        'is_finalized': False,
                        'uploaded_file': excel_file,
                        'file_type': 'excel',
                    }
                )
                
                logger.info(f"DEBUG: Period {'created' if created else 'updated'} - ID: {period.id}")
                
                # If updating existing period, save the new excel file
                if not created:
                    period.uploaded_file = excel_file
                    period.file_type = 'excel'
                    period.save()
                    logger.info(f"DEBUG: Updated uploaded_file for period {period.id}")
                
                logger.info(f"DEBUG: Saving trading account data")
                TradingAccount.objects.update_or_create(
                    period=period,
                    defaults=trading_account_data
                )
                
                logger.info(f"DEBUG: Saving profit & loss data")
                # Ensure defaults are applied before saving (safety check)
                profit_loss_data = self._default_profit_loss(profit_loss_data)
                logger.info(f"DEBUG: Profit & Loss data before save: {profit_loss_data}")
                logger.info(f"DEBUG: miscellaneous_income value: {profit_loss_data.get('miscellaneous_income')}")
                
                # Final safety check - ensure miscellaneous_income is never null
                if 'miscellaneous_income' not in profit_loss_data or profit_loss_data.get('miscellaneous_income') is None:
                    profit_loss_data['miscellaneous_income'] = Decimal('0')
                    logger.warning(f"DEBUG: ⚠️ miscellaneous_income was null, set to 0")
                
                # Create/Update Profit & Loss
                ProfitAndLoss.objects.update_or_create(
                    period=period,
                    defaults=profit_loss_data
                )
                
                logger.info(f"DEBUG: Saving balance sheet data")
                # Create/Update Balance Sheet
                BalanceSheet.objects.update_or_create(
                    period=period,
                    defaults=balance_sheet_data
                )
                
                logger.info(f"DEBUG: Saving operational metrics data")
                # Ensure staff_count is always set before saving (safety check)
                if 'staff_count' not in operational_metrics_data or operational_metrics_data.get('staff_count') is None:
                    operational_metrics_data['staff_count'] = 1
                    logger.warning(f"DEBUG: staff_count was missing or null, setting to default: 1")
                
                logger.info(f"DEBUG: Operational metrics data before save: {operational_metrics_data}")
                
                # Create/Update Operational Metrics
                OperationalMetrics.objects.update_or_create(
                    period=period,
                    defaults=operational_metrics_data
                )
                
                logger.info(f"DEBUG: Calculating ratios")
                # Automatically calculate ratios
                from app.services.ratio_calculator import RatioCalculator
                calculator = RatioCalculator(period)
                all_ratios = calculator.calculate_all_ratios()
                traffic_light_statuses = calculator.get_traffic_light_statuses()
                
                logger.info(f"DEBUG: Saving ratio results")
                # Create or update RatioResult
                RatioResult.objects.update_or_create(
                    period=period,
                    defaults={
                        'working_fund': Decimal(str(all_ratios['working_fund'])),
                        'stock_turnover': Decimal(str(all_ratios.get('stock_turnover', 0))),
                        'gross_profit_ratio': Decimal(str(all_ratios.get('gross_profit_ratio', 0))),
                        'net_profit_ratio': Decimal(str(all_ratios.get('net_profit_ratio', 0))),
                        'own_fund_to_wf': Decimal(str(all_ratios.get('own_fund_to_wf', 0))),
                        'deposits_to_wf': Decimal(str(all_ratios.get('deposits_to_wf', 0))),
                        'borrowings_to_wf': Decimal(str(all_ratios.get('borrowings_to_wf', 0))),
                        'loans_to_wf': Decimal(str(all_ratios.get('loans_to_wf', 0))),
                        'investments_to_wf': Decimal(str(all_ratios.get('investments_to_wf', 0))),
                        'cost_of_deposits': Decimal(str(all_ratios.get('cost_of_deposits', 0))),
                        'yield_on_loans': Decimal(str(all_ratios.get('yield_on_loans', 0))),
                        'yield_on_investments': Decimal(str(all_ratios.get('yield_on_investments', 0))),
                        'credit_deposit_ratio': Decimal(str(all_ratios.get('credit_deposit_ratio', 0))),
                        'avg_cost_of_wf': Decimal(str(all_ratios.get('avg_cost_of_wf', 0))),
                        'avg_yield_on_wf': Decimal(str(all_ratios.get('avg_yield_on_wf', 0))),
                        'gross_fin_margin': Decimal(str(all_ratios.get('gross_fin_margin', 0))),
                        'operating_cost_to_wf': Decimal(str(all_ratios.get('operating_cost_to_wf', 0))),
                        'net_fin_margin': Decimal(str(all_ratios.get('net_fin_margin', 0))),
                        'risk_cost_to_wf': Decimal(str(all_ratios.get('risk_cost_to_wf', 0))),
                        'net_margin': Decimal(str(all_ratios.get('net_margin', 0))),
                        'all_ratios': all_ratios,
                        'traffic_light_status': traffic_light_statuses
                    }
                )
                
                logger.info(f"DEBUG: All data saved successfully for period {period.id}")
            
            logger.info(f"DEBUG: Returning success response")
            return Response({
                "status": "success",
                "response_code": status.HTTP_200_OK,
                "message": "Excel file processed successfully",
                "period_id": period.id,
                "period_label": period.label
            })
            
            
        except Exception as e:
            logger.exception(f"=== EXCEPTION in UploadExcelView: {str(e)} ===")
            import traceback
            traceback.print_exc()
            return Response({
                "status": "failed",
                "response_code": status.HTTP_500_INTERNAL_SERVER_ERROR,
                "message": f"Error processing Excel file: {str(e)}"
            })
    
    def _parse_balance_sheet(self, sheet, company=None):
        """Parse Balance Sheet sheet - handles both column format (Liabilities/Amount, Assets/Amount) and row format"""
        logger.info(f"=== PARSING BALANCE SHEET ===")
        logger.info(f"Sheet name: {sheet.title}")
        
        # Print all rows from the sheet
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(row_data)
            logger.info(f"Row {row_idx}: {row_data}")
        
        logger.info(f"Total rows in Balance Sheet: {len(all_rows)}")
        
        data = {}
        
        # Try to detect format by checking first row
        first_row = [str(cell).strip().lower() if cell else '' for cell in (all_rows[0] if len(all_rows) > 0 else [])]
        logger.info(f"First row (headers): {first_row}")
        
        # Check if it's column format (Liabilities/Amount/Assets/Amount)
        if 'liabilities' in ' '.join(first_row) or 'amount' in ' '.join(first_row):
            # Column format: Liabilities | Amount | Assets | Amount
            liabilities_col = None
            liabilities_amount_col = None
            assets_col = None
            assets_amount_col = None
            
            for row_idx, row in enumerate(all_rows, 1):
                row_values = [cell if cell else None for cell in row]
                logger.info(f"Processing Balance Sheet Row {row_idx}: {row_values}")
                
                if row_idx == 1:
                    # Find column indices
                    for col_idx, cell_value in enumerate(row_values):
                        if cell_value:
                            header = str(cell_value).strip().lower()
                            if 'liabilities' in header:
                                liabilities_col = col_idx
                                logger.info(f"Found Liabilities column at index {col_idx}")
                            elif 'assets' in header and liabilities_col is not None:
                                assets_col = col_idx
                                logger.info(f"Found Assets column at index {col_idx}")
                            elif 'amount' in header:
                                if liabilities_amount_col is None:
                                    liabilities_amount_col = col_idx
                                    logger.info(f"Found Liabilities Amount column at index {col_idx}")
                                else:
                                    assets_amount_col = col_idx
                                    logger.info(f"Found Assets Amount column at index {col_idx}")
                elif row_idx > 1:
                    # Parse data rows
                    if liabilities_col is not None and liabilities_amount_col is not None:
                        item = row_values[liabilities_col] if liabilities_col < len(row_values) else None
                        amount = row_values[liabilities_amount_col] if liabilities_amount_col < len(row_values) else None
                        logger.info(f"Row {row_idx} - Liability: item='{item}', amount='{amount}'")
                        if item and amount is not None:
                            field_name = self._map_balance_sheet_field(str(item), company)
                            if field_name:
                                data[field_name] = self._parse_decimal(amount)
                                logger.info(f"✓ Row {row_idx}: Mapped liability '{item}' -> {field_name} = {amount}")
                            else:
                                logger.warning(f"✗ Row {row_idx}: Could not map liability field: '{item}'")
                    
                    if assets_col is not None and assets_amount_col is not None:
                        item = row_values[assets_col] if assets_col < len(row_values) else None
                        amount = row_values[assets_amount_col] if assets_amount_col < len(row_values) else None
                        logger.info(f"Row {row_idx} - Asset: item='{item}', amount='{amount}'")
                        if item and amount is not None:
                            field_name = self._map_balance_sheet_field(str(item), company)
                            if field_name:
                                data[field_name] = self._parse_decimal(amount)
                                logger.info(f"✓ Row {row_idx}: Mapped asset '{item}' -> {field_name} = {amount}")
                            else:
                                logger.warning(f"✗ Row {row_idx}: Could not map asset field: '{item}'")
        else:
            # Row format: Headers in first row, data in second row
            logger.info("DEBUG: Using row format (headers in row 1, data in row 2)")
            headers = {}
            data_row = None
            
            for row_idx, row in enumerate(all_rows, 1):
                row_values = [cell if cell else None for cell in row]
                logger.info(f"Balance Sheet Row format - Row {row_idx}: {row_values}")
                
                if row_idx == 1:
                    for col_idx, cell_value in enumerate(row_values):
                        if cell_value:
                            header = str(cell_value).strip().lower()
                            headers[header] = col_idx
                            logger.info(f"Found header '{header}' at column {col_idx}")
                elif row_idx == 2:
                    data_row = row_values
                    logger.info(f"Row 2 (Data row): {data_row}")
                    break
            
            if data_row and headers:
                logger.info("Parsing balance sheet data row...")
                for header, col_idx in headers.items():
                    if col_idx < len(data_row):
                        value = data_row[col_idx]
                        logger.info(f"Balance Sheet - Header: '{header}', Column: {col_idx}, Value: '{value}'")
                        if value is not None:
                            field_name = self._map_balance_sheet_field(header, company)
                            if field_name:
                                data[field_name] = self._parse_decimal(value)
                                logger.info(f"✓ Mapped '{header}' -> {field_name} = {value}")
                            else:
                                logger.warning(f"✗ Could not map header: '{header}'")
        
        logger.info(f"=== BALANCE SHEET PARSING COMPLETE ===")
        logger.info(f"Extracted data: {data}")
        return data
    
    def _map_balance_sheet_field(self, item_str, company=None):
        """Map balance sheet item string to model field name. Uses StatementColumnConfig (checks company-specific first, then global, including aliases)."""
        # Check StatementColumnConfig (checks company-specific first, then global, including aliases)
        normalized = StatementColumnConfig._normalize_for_match(item_str)
        logger.debug(f"Mapping Balance Sheet field: '{item_str}' (normalized: '{normalized}') for company: {company.name if company else 'None'}")
        
        canonical = StatementColumnConfig.resolve_canonical_field(company, "BALANCE_SHEET", item_str)
        if canonical:
            logger.debug(f"✓ Found mapping: '{item_str}' -> {canonical}")
            return canonical
        
        # Fallback: Try pattern matching against common model field names (if StatementColumnConfig doesn't have entry)
        # Clean normalized string for better matching (remove special chars, keep only alphanumeric)
        normalized_clean = normalized.replace('&', '').replace('/', '').replace('(', '').replace(')', '').replace('-', '').replace(' ', '').replace('_', '').lower()
        
        # Pattern matching for common variations
        if 'share' in normalized_clean and 'capital' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> share_capital")
            return 'share_capital'
        if ('member' in normalized_clean or 'deposit' in normalized_clean) and 'deposit' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> deposits")
            return 'deposits'
        if 'provision' in normalized_clean and 'made' not in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> provisions")
            return 'provisions'
        if 'investment' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> investments")
            return 'investments'
        if 'other' in normalized_clean and 'asset' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> other_assets")
            return 'other_assets'
        if ('stock' in normalized_clean or 'trade' in normalized_clean) and 'stock' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> stock_in_trade")
            return 'stock_in_trade'
        if 'loan' in normalized_clean and 'advance' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> loans_advances")
            return 'loans_advances'
        if 'reserve' in normalized_clean and ('statutory' in normalized_clean or 'free' in normalized_clean):
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> reserves_statutory_free")
            return 'reserves_statutory_free'
        if 'statutory' in normalized_clean and 'free' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> reserves_statutory_free")
            return 'reserves_statutory_free'
        
        # Direct match if normalized name matches exactly
        common_balance_sheet_fields = [
            'share_capital', 'deposits', 'borrowings', 'reserves_statutory_free', 
            'undistributed_profit', 'provisions', 'other_liabilities', 
            'cash_in_hand', 'cash_at_bank', 'investments', 'loans_advances', 
            'fixed_assets', 'other_assets', 'stock_in_trade'
        ]
        
        if normalized in common_balance_sheet_fields:
            logger.info(f"✓ Fallback: Direct match '{item_str}' (normalized: '{normalized}') -> {normalized}")
            return normalized
        
        logger.warning(f"✗ No mapping found in StatementColumnConfig for Balance Sheet field: '{item_str}' (normalized: '{normalized}')")
        return None
    
    def _parse_profit_loss(self, sheet, company=None):
        """Parse Profit and Loss sheet - handles Expenses/Amount and Income/Amount column format"""
        logger.info(f"=== PARSING PROFIT & LOSS SHEET ===")
        logger.info(f"Sheet name: {sheet.title}")
        
        # Print all rows from the sheet
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(row_data)
            logger.info(f"Row {row_idx}: {row_data}")
        
        logger.info(f"Total rows in sheet: {len(all_rows)}")
        
        data = {}
        
        # Check format - column format (Expenses/Amount, Income/Amount) or row format
        first_row = [str(cell).strip().lower() if cell else '' for cell in (all_rows[0] if len(all_rows) > 0 else [])]
        logger.info(f"First row (headers): {first_row}")
        
        if 'expenses' in ' '.join(first_row) or 'income' in ' '.join(first_row):
            # Column format: Expenses | Amount | Income | Amount
            expenses_col = None
            expenses_amount_col = None
            income_col = None
            income_amount_col = None
            
            for row_idx, row in enumerate(all_rows, 1):
                row_values = [cell if cell else None for cell in row]
                logger.info(f"Processing Row {row_idx}: {row_values}")
                
                if row_idx == 1:
                    for col_idx, cell_value in enumerate(row_values):
                        if cell_value:
                            header = str(cell_value).strip().lower()
                            if 'expenses' in header:
                                expenses_col = col_idx
                                logger.info(f"Found Expenses column at index {col_idx}")
                            elif 'income' in header:
                                income_col = col_idx
                                logger.info(f"Found Income column at index {col_idx}")
                            elif 'amount' in header:
                                if expenses_amount_col is None:
                                    expenses_amount_col = col_idx
                                    logger.info(f"Found Expenses Amount column at index {col_idx}")
                                else:
                                    income_amount_col = col_idx
                                    logger.info(f"Found Income Amount column at index {col_idx}")
                elif row_idx > 1:
                    # Parse expenses
                    if expenses_col is not None and expenses_amount_col is not None:
                        item = row_values[expenses_col] if expenses_col < len(row_values) else None
                        amount = row_values[expenses_amount_col] if expenses_amount_col < len(row_values) else None
                        logger.info(f"Row {row_idx} - Expense: item='{item}', amount='{amount}'")
                        if item and amount is not None:
                            item_str = str(item).strip()
                            field_name = self._map_profit_loss_field(item_str, is_income=False, company=company)
                            if field_name:
                                data[field_name] = self._parse_decimal(amount)
                                logger.info(f"✓ Row {row_idx}: Mapped expense '{item_str}' -> {field_name} = {amount}")
                            else:
                                logger.warning(f"✗ Row {row_idx}: Could not map expense field: '{item_str}'")
                    
                    # Parse income
                    if income_col is not None and income_amount_col is not None:
                        item = row_values[income_col] if income_col < len(row_values) else None
                        amount = row_values[income_amount_col] if income_amount_col < len(row_values) else None
                        logger.info(f"Row {row_idx} - Income: item='{item}', amount='{amount}'")
                        if item and amount is not None:
                            item_str = str(item).strip()
                            field_name = self._map_profit_loss_field(item_str, is_income=True, company=company)
                            if field_name:
                                data[field_name] = self._parse_decimal(amount)
                                logger.info(f"✓ Row {row_idx}: Mapped income '{item_str}' -> {field_name} = {amount}")
                            else:
                                logger.warning(f"✗ Row {row_idx}: Could not map income field: '{item_str}'")
        else:
            # Row format: Headers in first row, income in row 2, expenses in row 3
            logger.info("DEBUG: Using row format (headers in row 1, income in row 2, expenses in row 3)")
            headers = {}
            income_row = None
            expense_row = None
            
            for row_idx, row in enumerate(all_rows, 1):
                row_values = [cell if cell else None for cell in row]
                logger.info(f"Row format - Row {row_idx}: {row_values}")
                
                if row_idx == 1:
                    for col_idx, cell_value in enumerate(row_values):
                        if cell_value:
                            header = str(cell_value).strip().lower()
                            headers[header] = col_idx
                            logger.info(f"Found header '{header}' at column {col_idx}")
                elif row_idx == 2:
                    income_row = row_values
                    logger.info(f"Row 2 (Income row): {income_row}")
                elif row_idx == 3:
                    expense_row = row_values
                    logger.info(f"Row 3 (Expense row): {expense_row}")
                    break
            
            if headers:
                # Parse income
                if income_row:
                    logger.info("Parsing income row...")
                    for header, col_idx in headers.items():
                        if col_idx < len(income_row):
                            value = income_row[col_idx]
                            logger.info(f"Income - Header: '{header}', Column: {col_idx}, Value: '{value}'")
                            if value is not None:
                                field_name = self._map_profit_loss_field(header, is_income=True, company=company)
                                if field_name:
                                    data[field_name] = self._parse_decimal(value)
                                    logger.info(f"✓ Mapped income '{header}' -> {field_name} = {value}")
                                else:
                                    logger.warning(f"✗ Could not map income header: '{header}'")
                
                # Parse expenses
                if expense_row:
                    logger.info("Parsing expense row...")
                    for header, col_idx in headers.items():
                        if col_idx < len(expense_row):
                            value = expense_row[col_idx]
                            logger.info(f"Expense - Header: '{header}', Column: {col_idx}, Value: '{value}'")
                            if value is not None:
                                field_name = self._map_profit_loss_field(header, is_income=False, company=company)
                                if field_name:
                                    data[field_name] = self._parse_decimal(value)
                                    logger.info(f"✓ Mapped expense '{header}' -> {field_name} = {value}")
                                else:
                                    logger.warning(f"✗ Could not map expense header: '{header}'")
        
        logger.info(f"=== PROFIT & LOSS PARSING COMPLETE ===")
        logger.info(f"Extracted data: {data}")
        return data
    
    def _map_profit_loss_field(self, item_str, is_income=False, company=None):
        """Map profit & loss item string to model field name. Uses StatementColumnConfig (checks company-specific first, then global, including aliases)."""
        # Check StatementColumnConfig (checks company-specific first, then global, including aliases)
        normalized = StatementColumnConfig._normalize_for_match(item_str)
        logger.debug(f"Mapping Profit & Loss field: '{item_str}' (normalized: '{normalized}', is_income={is_income}) for company: {company.name if company else 'None'}")
        
        canonical = StatementColumnConfig.resolve_canonical_field(company, "PL", item_str)
        if canonical:
            logger.debug(f"✓ Found mapping: '{item_str}' -> {canonical}")
            return canonical
        
        # Fallback: Try pattern matching against common model field names (if StatementColumnConfig doesn't have entry)
        # Clean normalized string for better matching (remove special chars, handle a/c -> ac)
        normalized_clean = normalized.replace('&', '').replace('/', '').replace('(', '').replace(')', '').replace('-', '').replace(' ', '').replace('_', '').replace('a/c', 'ac').replace('ac', 'ac').lower()
        
        # Pattern matching for common variations
        if 'interest' in normalized_clean and 'deposit' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> interest_on_deposits")
            return 'interest_on_deposits'
        if 'interest' in normalized_clean and 'bank' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> interest_on_bank_ac")
            return 'interest_on_bank_ac'
        if ('establishment' in normalized_clean or 'establishm' in normalized_clean) and ('contingenc' in normalized_clean or 'conting' in normalized_clean):
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> establishment_contingencies")
            return 'establishment_contingencies'
        if 'provision' in normalized_clean and 'made' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> provisions")
            return 'provisions'
        if 'net' in normalized_clean and 'profit' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> net_profit")
            return 'net_profit'
        if 'miscellaneous' in normalized_clean or 'miscellane' in normalized_clean:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> miscellaneous_income")
            return 'miscellaneous_income'
        
        # Direct match if normalized name matches exactly
        common_pl_fields = [
            'interest_on_loans', 'interest_on_bank_ac', 'return_on_investment', 
            'miscellaneous_income', 'interest_on_deposits', 'interest_on_borrowings', 
            'establishment_contingencies', 'provisions', 'net_profit'
        ]
        
        if normalized in common_pl_fields:
            logger.info(f"✓ Fallback: Direct match '{item_str}' (normalized: '{normalized}') -> {normalized}")
            return normalized
        
        logger.warning(f"✗ No mapping found in StatementColumnConfig for Profit & Loss field: '{item_str}' (normalized: '{normalized}')")
        return None
    
    def _parse_trading_account(self, sheet, company=None):
        """Parse Trading Account sheet. Uses StatementColumnConfig (display_name, aliases) for column name matching first."""
        logger.info(f"=== PARSING TRADING ACCOUNT SHEET ===")
        logger.info(f"Sheet name: {sheet.title}")
        
        # Print all rows from the sheet
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(row_data)
            logger.info(f"Row {row_idx}: {row_data}")
        
        logger.info(f"Total rows in Trading Account: {len(all_rows)}")
        
        data = {}
        
        # Look for Item and Amount columns
        headers = {}
        item_col = None
        amount_col = None
        
        for row_idx, row in enumerate(all_rows, 1):
            row_values = [cell if cell else None for cell in row]
            logger.info(f"Processing Trading Account Row {row_idx}: {row_values}")
            if row_idx == 1:
                for col_idx, cell_value in enumerate(row_values):
                    if cell_value:
                        header = str(cell_value).strip().lower()
                        if 'item' in header:
                            item_col = col_idx
                            logger.info(f"Found Item column at index {col_idx}")
                        elif 'amount' in header:
                            amount_col = col_idx
                            logger.info(f"Found Amount column at index {col_idx}")
            elif row_idx > 1 and item_col is not None and amount_col is not None:
                item = row_values[item_col] if item_col < len(row_values) else None
                amount = row_values[amount_col] if amount_col < len(row_values) else None
                
                logger.info(f"Row {row_idx} - Trading Account: item='{item}', amount='{amount}'")
                if item and amount is not None:
                    item_str = str(item).strip()
                    field_name = self._map_trading_account_field(item_str, company)
                    if field_name:
                        data[field_name] = self._parse_decimal(amount)
                        logger.info(f"✓ Row {row_idx}: Mapped '{item_str}' -> {field_name} = {amount}")
                    else:
                        logger.warning(f"✗ Row {row_idx}: Could not map trading account field: '{item_str}'")
        
        logger.info(f"=== TRADING ACCOUNT PARSING COMPLETE ===")
        logger.info(f"Extracted data: {data}")
        return data
    
    def _map_trading_account_field(self, item_str, company=None):
        """Map trading account item string to canonical field. Uses StatementColumnConfig (checks company-specific first, then global, including aliases)."""
        # Check StatementColumnConfig (checks company-specific first, then global, including aliases)
        normalized = StatementColumnConfig._normalize_for_match(item_str)
        logger.debug(f"Mapping Trading Account field: '{item_str}' (normalized: '{normalized}') for company: {company.name if company else 'None'}")
        
        canonical = StatementColumnConfig.resolve_canonical_field(company, "TRADING", item_str)
        if canonical:
            logger.debug(f"✓ Found mapping: '{item_str}' -> {canonical}")
            return canonical
        
        # Fallback: Try direct match against common model field names (if StatementColumnConfig doesn't have entry)
        common_trading_fields = [
            'opening_stock', 'purchases', 'trade_charges', 'sales', 'closing_stock'
        ]
        
        # Also check for common variations
        if normalized in common_trading_fields:
            logger.info(f"✓ Fallback: Direct match '{item_str}' (normalized: '{normalized}') -> {normalized}")
            return normalized
        
        # Handle common variations
        if 'opening' in normalized and ('stock' in normalized or 'inventory' in normalized):
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> opening_stock")
            return 'opening_stock'
        if 'closing' in normalized and ('stock' in normalized or 'inventory' in normalized):
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> closing_stock")
            return 'closing_stock'
        if 'purchase' in normalized:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> purchases")
            return 'purchases'
        if 'sale' in normalized:
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> sales")
            return 'sales'
        if ('trade' in normalized or 'direct' in normalized) and ('charge' in normalized or 'expense' in normalized):
            logger.info(f"✓ Fallback: Pattern match '{item_str}' -> trade_charges")
            return 'trade_charges'
        
        logger.warning(f"✗ No mapping found in StatementColumnConfig for Trading Account field: '{item_str}' (normalized: '{normalized}')")
        return None
    
    def _parse_operational_metrics(self, sheet, company=None):
        """Parse Operational Metrics sheet. Uses StatementColumnConfig for metric name matching when available."""
        logger.info(f"=== PARSING OPERATIONAL METRICS SHEET ===")
        logger.info(f"Sheet name: {sheet.title}")
        
        # Print all rows from the sheet
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(row_data)
            logger.info(f"Row {row_idx}: {row_data}")
        
        logger.info(f"Total rows in Operational Metrics: {len(all_rows)}")
        
        data = {}
        
        # Look for Metric and Value columns
        metric_col = None
        value_col = None
        
        for row_idx, row in enumerate(all_rows, 1):
            row_values = [cell if cell else None for cell in row]
            logger.info(f"Processing Operational Metrics Row {row_idx}: {row_values}")
            if row_idx == 1:
                for col_idx, cell_value in enumerate(row_values):
                    if cell_value:
                        header = str(cell_value).strip().lower()
                        if 'metric' in header:
                            metric_col = col_idx
                            logger.info(f"Found Metric column at index {col_idx}")
                        elif 'value' in header:
                            value_col = col_idx
                            logger.info(f"Found Value column at index {col_idx}")
            elif row_idx > 1 and metric_col is not None and value_col is not None:
                metric = row_values[metric_col] if metric_col < len(row_values) else None
                value = row_values[value_col] if value_col < len(row_values) else None
                logger.info(f"Row {row_idx} - Operational Metrics: metric='{metric}', value='{value}'")
                
                if metric and value is not None:
                    metric_str = str(metric).strip()
                    metric_lower = metric_str.lower()
                    field_name = None
                    
                    # First: Check direct model field mappings
                    if 'staff' in metric_lower and 'count' in metric_lower:
                        field_name = 'staff_count'
                    
                    # Second: Fallback to StatementColumnConfig (for custom/company-specific mappings)
                    if not field_name and company:
                        canonical = StatementColumnConfig.resolve_canonical_field(company, "OPERATIONAL", metric_str)
                        if canonical:
                            field_name = canonical
                    
                    if field_name:
                        try:
                            data[field_name] = int(float(value))
                            logger.info(f"✓ Row {row_idx}: Mapped '{metric_str}' -> {field_name} = {data[field_name]}")
                        except (TypeError, ValueError):
                            if field_name == 'staff_count':
                                data[field_name] = 1
                                logger.warning(f"✗ Row {row_idx}: Could not parse value '{value}' for '{metric_str}', defaulting to 1")
                            else:
                                data[field_name] = value
                                logger.warning(f"✗ Row {row_idx}: Could not parse value '{value}' for '{metric_str}', using raw value")
                    else:
                        logger.warning(f"✗ Row {row_idx}: Could not map operational metric: '{metric_str}'")
        
        # Ensure staff_count is always set (default to 1 if not found)
        if 'staff_count' not in data:
            data['staff_count'] = 1
            logger.info("Set default staff_count = 1 (not found in sheet)")
        
        logger.info(f"=== OPERATIONAL METRICS PARSING COMPLETE ===")
        logger.info(f"Extracted data: {data}")
        return data
    
    def _parse_financial_statement_sheet(self, sheet):
        """Parse Financial_Statement sheet: Entity Name, Fiscal Year End, Currency, Staff Count"""
        data = {}
        col_map = {}  # header -> col index
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            if row_idx == 1:
                for col_idx, cell_value in enumerate(row):
                    if cell_value:
                        header = str(cell_value).strip().lower().replace(' ', '_')
                        col_map[header] = col_idx
            else:
                if not col_map:
                    break
                for header, col_idx in col_map.items():
                    if col_idx >= len(row):
                        continue
                    val = row[col_idx]
                    if val is None:
                        continue
                    if 'entity' in header and 'name' in header:
                        data['entity_name'] = str(val).strip()
                    elif 'fiscal' in header and 'year' in header:
                        data['fiscal_year_end'] = str(val).strip()
                    elif 'currency' in header:
                        data['currency'] = str(val).strip()
                    elif 'staff' in header and 'count' in header:
                        try:
                            data['staff_count'] = int(float(val))
                        except (TypeError, ValueError):
                            pass
                break
        return data
    
    def _parse_balance_sheet_liabilities(self, sheet, company=None):
        """Parse Balance_Sheet_Liabilities: Liability Type, Amount (one row per liability). Uses StatementColumnConfig first."""
        logger.info(f"=== PARSING BALANCE SHEET LIABILITIES ===")
        logger.info(f"Sheet name: {sheet.title}")
        
        # Print all rows from the sheet
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(row_data)
            logger.info(f"Row {row_idx}: {row_data}")
        
        logger.info(f"Total rows in Balance Sheet Liabilities: {len(all_rows)}")
        
        data = {}
        type_col = amount_col = None
        liability_to_field = [
            ('reserves (statutory + free)', 'reserves_statutory_free'),
            ('reserves (statutory', 'reserves_statutory_free'),
            ('share capital', 'share_capital'),
            ('deposits', 'deposits'),
            ('borrowings', 'borrowings'),
            ('reserves', 'reserves_statutory_free'),
            ('statutory', 'reserves_statutory_free'),
            ('provisions', 'provisions'),
            ('other liabilities', 'other_liabilities'),
            ('undistributed profit', 'undistributed_profit'),
            ('udp', 'undistributed_profit'),
        ]
        for row_idx, row in enumerate(all_rows, 1):
            row_values = [cell if cell else None for cell in row]
            logger.info(f"Processing Liabilities Row {row_idx}: {row_values}")
            
            if row_idx == 1:
                for col_idx, cell_value in enumerate(row_values):
                    if cell_value:
                        h = str(cell_value).strip().lower()
                        if 'liability' in h or 'type' in h:
                            type_col = col_idx
                            logger.info(f"Found Liability Type column at index {col_idx}")
                        elif 'amount' in h:
                            amount_col = col_idx
                            logger.info(f"Found Amount column at index {col_idx}")
            elif row_idx > 1 and type_col is not None and amount_col is not None:
                typ = row_values[type_col] if type_col < len(row_values) else None
                amt = row_values[amount_col] if amount_col < len(row_values) else None
                logger.info(f"Row {row_idx} - Liability: type='{typ}', amount='{amt}'")
                if typ is not None and amt is not None:
                    t = str(typ).strip()
                    field_name = None
                    # First: Check direct model field mappings
                    t_lower = t.lower()
                    for key, field in liability_to_field:
                        if key in t_lower or t_lower in key:
                            field_name = field
                            break
                    # Second: Fallback to StatementColumnConfig (for custom/company-specific mappings)
                    if not field_name and company:
                        field_name = StatementColumnConfig.resolve_canonical_field(company, "BALANCE_SHEET", t)
                    if field_name:
                        data[field_name] = self._parse_decimal(amt)
                        logger.info(f"✓ Row {row_idx}: Mapped liability '{t}' -> {field_name} = {amt}")
                    else:
                        logger.warning(f"✗ Row {row_idx}: Could not map liability type: '{t}'")
        
        logger.info(f"=== BALANCE SHEET LIABILITIES PARSING COMPLETE ===")
        logger.info(f"Extracted data: {data}")
        return data
    
    def _parse_balance_sheet_assets(self, sheet, company=None):
        """Parse Balance_Sheet_Assets: Asset Type, Amount (one row per asset). Uses StatementColumnConfig first."""
        logger.info(f"=== PARSING BALANCE SHEET ASSETS ===")
        logger.info(f"Sheet name: {sheet.title}")
        
        # Print all rows from the sheet
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(row_data)
            logger.info(f"Row {row_idx}: {row_data}")
        
        logger.info(f"Total rows in Balance Sheet Assets: {len(all_rows)}")
        
        data = {}
        type_col = amount_col = None
        asset_to_field = {
            'cash in hand': 'cash_in_hand',
            'cash at bank': 'cash_at_bank',
            'investments': 'investments',
            'loans & advances': 'loans_advances',
            'loans and advances': 'loans_advances',
            'fixed assets': 'fixed_assets',
            'other assets': 'other_assets',
            'stock in trade': 'stock_in_trade',
        }
        for row_idx, row in enumerate(all_rows, 1):
            row_values = [cell if cell else None for cell in row]
            logger.info(f"Processing Assets Row {row_idx}: {row_values}")
            
            if row_idx == 1:
                for col_idx, cell_value in enumerate(row_values):
                    if cell_value:
                        h = str(cell_value).strip().lower()
                        if 'asset' in h or 'type' in h:
                            type_col = col_idx
                            logger.info(f"Found Asset Type column at index {col_idx}")
                        elif 'amount' in h:
                            amount_col = col_idx
                            logger.info(f"Found Amount column at index {col_idx}")
            elif row_idx > 1 and type_col is not None and amount_col is not None:
                typ = row_values[type_col] if type_col < len(row_values) else None
                amt = row_values[amount_col] if amount_col < len(row_values) else None
                logger.info(f"Row {row_idx} - Asset: type='{typ}', amount='{amt}'")
                if typ is not None and amt is not None:
                    t = str(typ).strip()
                    field_name = None
                    # First: Check direct model field mappings
                    t_lower = t.lower()
                    for key, field in asset_to_field.items():
                        if key in t_lower or t_lower in key:
                            field_name = field
                            break
                    # Second: Fallback to StatementColumnConfig (for custom/company-specific mappings)
                    if not field_name and company:
                        field_name = StatementColumnConfig.resolve_canonical_field(company, "BALANCE_SHEET", t)
                    if field_name:
                        data[field_name] = self._parse_decimal(amt)
                        logger.info(f"✓ Row {row_idx}: Mapped asset '{t}' -> {field_name} = {amt}")
                    else:
                        logger.warning(f"✗ Row {row_idx}: Could not map asset type: '{t}'")
        
        logger.info(f"=== BALANCE SHEET ASSETS PARSING COMPLETE ===")
        logger.info(f"Extracted data: {data}")
        return data
    
    def _parse_profit_loss_rows(self, sheet, company=None):
        """Parse Profit_Loss sheet: Category, Item, Amount. Uses StatementColumnConfig for item name matching first."""
        logger.info(f"=== PARSING PROFIT & LOSS ROWS (Format A) ===")
        logger.info(f"Sheet name: {sheet.title}")
        
        # Print all rows from the sheet
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(row_data)
            logger.info(f"Row {row_idx}: {row_data}")
        
        logger.info(f"Total rows in Profit & Loss: {len(all_rows)}")
        
        data = {}
        cat_col = item_col = amount_col = None
        for row_idx, row in enumerate(all_rows, 1):
            row_values = [cell if cell else None for cell in row]
            logger.info(f"Processing Profit & Loss Row {row_idx}: {row_values}")
            if row_idx == 1:
                for col_idx, cell_value in enumerate(row_values):
                    if cell_value:
                        h = str(cell_value).strip().lower()
                        if 'category' in h:
                            cat_col = col_idx
                            logger.info(f"Found Category column at index {col_idx}")
                        elif 'item' in h:
                            item_col = col_idx
                            logger.info(f"Found Item column at index {col_idx}")
                        elif 'amount' in h:
                            amount_col = col_idx
                            logger.info(f"Found Amount column at index {col_idx}")
            elif row_idx > 1 and item_col is not None and amount_col is not None:
                cat = str(row_values[cat_col] or '').strip().lower() if cat_col is not None and cat_col < len(row_values) else ''
                item = str(row_values[item_col] or '').strip() if item_col < len(row_values) else ''
                amt = row_values[amount_col] if amount_col < len(row_values) else None
                
                logger.info(f"Row {row_idx} - Profit & Loss: category='{cat}', item='{item}', amount='{amt}'")
                
                if not item or amt is None:
                    logger.debug(f"Row {row_idx}: Skipping (empty item or amount)")
                    continue
                amt_val = self._parse_decimal(amt)
                item_lower = item.lower()
                field_name = None
                
                # First: Check direct model field mappings based on category and item
                if 'income' in cat:
                    if 'interest' in item_lower and 'loan' in item_lower:
                        field_name = 'interest_on_loans'
                    elif 'interest' in item_lower and 'bank' in item_lower:
                        field_name = 'interest_on_bank_ac'
                    elif 'return' in item_lower and 'investment' in item_lower:
                        field_name = 'return_on_investment'
                    elif 'miscellaneous' in item_lower or 'miscellane' in item_lower:
                        field_name = 'miscellaneous_income'
                elif 'expense' in cat:
                    if 'interest' in item_lower and 'deposit' in item_lower:
                        field_name = 'interest_on_deposits'
                    elif 'interest' in item_lower and 'borrowing' in item_lower:
                        field_name = 'interest_on_borrowings'
                    elif 'establishment' in item_lower or 'contingenc' in item_lower or 'establishm' in item_lower:
                        field_name = 'establishment_contingencies'
                    elif 'provision' in item_lower:
                        field_name = 'provisions'
                elif 'net profit' in cat or (cat == '' and 'net profit' in item_lower):
                    field_name = 'net_profit'
                
                # Second: Fallback to StatementColumnConfig (for custom/company-specific mappings)
                if not field_name and company:
                    canonical = StatementColumnConfig.resolve_canonical_field(company, "PL", item)
                    if canonical:
                        field_name = canonical
                
                if field_name:
                    data[field_name] = amt_val
                    logger.info(f"✓ Row {row_idx}: Mapped '{item}' (category: {cat}) -> {field_name} = {amt_val}")
                else:
                    logger.warning(f"✗ Row {row_idx}: Could not map item '{item}' (category: {cat})")
        
        logger.info(f"=== PROFIT & LOSS ROWS PARSING COMPLETE ===")
        logger.info(f"Extracted data: {data}")
        return data
    
    def _parse_trading_account_rows(self, sheet, company=None):
        """Parse Trading_Account sheet: Item, Amount. Uses StatementColumnConfig for item name matching first."""
        logger.info(f"=== PARSING TRADING ACCOUNT ROWS (Format A) ===")
        logger.info(f"Sheet name: {sheet.title}")
        
        # Print all rows from the sheet
        all_rows = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            all_rows.append(row_data)
            logger.info(f"Row {row_idx}: {row_data}")
        
        logger.info(f"Total rows in Trading Account: {len(all_rows)}")
        
        data = {}
        item_col = amount_col = None
        for row_idx, row in enumerate(all_rows, 1):
            row_values = [cell if cell else None for cell in row]
            logger.info(f"Processing Trading Account Row {row_idx}: {row_values}")
            
            if row_idx == 1:
                for col_idx, cell_value in enumerate(row_values):
                    if cell_value:
                        h = str(cell_value).strip().lower()
                        if 'item' in h:
                            item_col = col_idx
                            logger.info(f"Found Item column at index {col_idx}")
                        elif 'amount' in h:
                            amount_col = col_idx
                            logger.info(f"Found Amount column at index {col_idx}")
            elif row_idx > 1 and item_col is not None and amount_col is not None:
                item = str(row_values[item_col] or '').strip() if item_col < len(row_values) else ''
                amt = row_values[amount_col] if amount_col < len(row_values) else None
                
                logger.info(f"Row {row_idx} - Trading Account: item='{item}', amount='{amt}'")
                
                if not item or amt is None:
                    logger.debug(f"Row {row_idx}: Skipping (empty item or amount)")
                    continue
                amt_val = self._parse_decimal(amt)
                field_name = self._map_trading_account_field(item, company)
                if field_name:
                    data[field_name] = amt_val
                    logger.info(f"✓ Row {row_idx}: Mapped '{item}' -> {field_name} = {amt_val}")
                else:
                    logger.warning(f"✗ Row {row_idx}: Could not map trading account item: '{item}'")
        
        logger.info(f"=== TRADING ACCOUNT ROWS PARSING COMPLETE ===")
        logger.info(f"Extracted data: {data}")
        return data
    
    def _parse_decimal(self, value):
        """Parse value to Decimal"""
        if value is None:
            return Decimal('0')
        if isinstance(value, (int, float)):
            return Decimal(str(value))
        if isinstance(value, str):
            clean_value = value.replace(',', '').strip()
            try:
                return Decimal(clean_value)
            except:
                return Decimal('0')
        return Decimal('0')
    
    def _default_balance_sheet(self, data):
        """Ensure all BalanceSheet fields exist (default 0)."""
        keys = [
            'share_capital', 'deposits', 'borrowings', 'reserves_statutory_free', 'undistributed_profit',
            'provisions', 'other_liabilities', 'cash_in_hand', 'cash_at_bank', 'investments',
            'loans_advances', 'fixed_assets', 'other_assets', 'stock_in_trade'
        ]
        for k in keys:
            if k not in data:
                data[k] = Decimal('0')
        return data
    
    def _default_profit_loss(self, data):
        """Ensure all ProfitAndLoss fields exist (default 0)."""
        keys = [
            'interest_on_loans', 'interest_on_bank_ac', 'return_on_investment', 'miscellaneous_income',
            'interest_on_deposits', 'interest_on_borrowings', 'establishment_contingencies', 'provisions',
            'net_profit'
        ]
        for k in keys:
            if k not in data or data.get(k) is None:
                data[k] = Decimal('0')
        return data
    
    def _default_trading_account(self, data):
        """Ensure all TradingAccount fields exist (default 0)."""
        keys = ['opening_stock', 'purchases', 'trade_charges', 'sales', 'closing_stock']
        for k in keys:
            if k not in data:
                data[k] = Decimal('0')
        return data
    
    def _find_sheets(self, workbook):
        """Find required sheets with flexible name matching. Supports two formats."""
        available_sheets = workbook.sheetnames
        sheet_mapping = {}
        
        # Format A: Financial_Statement, Balance_Sheet_Liabilities, Balance_Sheet_Assets, Profit_Loss, Trading_Account
        sheet_variations_a = {
            'Financial_Statement': ['financial_statement', 'financial statement'],
            'Balance_Sheet_Liabilities': ['balance_sheet_liabilities', 'balance sheet liabilities', 'liabilities'],
            'Balance_Sheet_Assets': ['balance_sheet_assets', 'balance sheet assets', 'assets'],
            'Profit_Loss': ['profit_loss', 'profit loss', 'profit_loss', 'profit & loss'],
            'Trading_Account': ['trading_account', 'trading account'],
        }
        # Format B: single Balance Sheet, P&L, Trading Account, Operational Metrics
        sheet_variations_b = {
            'Balance Sheet': ['balance sheet', 'balance_sheet', 'balancesheet', 'balance'],
            'Profit and Loss': ['profit and loss', 'profit & loss', 'profit_and_loss', 'profitandloss', 'p&l', 'pl'],
            'Trading Account': ['trading account', 'trading_account', 'tradingaccount', 'trading'],
            'Operational Metrics': ['operational metrics', 'operational_metrics', 'operationalmetrics', 'operational', 'metrics']
        }
        
        for required_name, variations in {**sheet_variations_a, **sheet_variations_b}.items():
            if required_name in sheet_mapping:
                continue
            for sheet_name in available_sheets:
                sheet_lower = sheet_name.lower().strip().replace(' ', '_')
                sheet_lower_orig = sheet_name.lower().strip()
                for variation in variations:
                    v_norm = variation.replace(' ', '_')
                    if v_norm in sheet_lower or sheet_lower == v_norm or variation in sheet_lower_orig or sheet_lower_orig == variation:
                        sheet_mapping[required_name] = sheet_name
                        break
                else:
                    continue
                break
        
        return sheet_mapping
    
    def _parse_docx_table(self, uploaded_file, company):
        """Parse .docx file with 4 tables format matching template: Balance Sheet, Profit & Loss, Trading Account, Operational Metrics."""
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            doc = Document(uploaded_file)
            
            # Get all tables and headings to identify each section
            tables = doc.tables
            if not tables:
                logger.warning("No tables found in .docx file")
                return {}, {}, {}, {}
            
            logger.info(f"DEBUG: Found {len(tables)} tables in .docx file")
            
            # Initialize data dicts
            balance_sheet_data = {}
            profit_loss_data = {}
            trading_account_data = {}
            operational_metrics_data = {}
            
            # Track which table we're processing (by index)
            table_index = 0
            
            # Parse each table based on its structure
            for table_idx, table in enumerate(tables):
                if len(table.rows) < 2:  # Need at least header + 1 data row
                    continue
                
                num_cols = len(table.rows[0].cells)
                logger.info(f"DEBUG: Table {table_idx + 1}: {len(table.rows)} rows, {num_cols} columns")
                
                # Identify table type by structure and content
                # Table 1: Balance Sheet - 4 columns (Liabilities, Amount, Assets, Amount)
                if num_cols == 4:
                    header_row = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    logger.info(f"DEBUG: Table {table_idx + 1} header: {header_row}")
                    
                    # Check if it's Balance Sheet or Profit & Loss by header
                    if 'liabilities' in ' '.join(header_row) and 'assets' in ' '.join(header_row):
                        # Balance Sheet table
                        logger.info(f"DEBUG: Parsing Balance Sheet table")
                        for row_idx in range(1, len(table.rows)):
                            row = table.rows[row_idx]
                            if len(row.cells) >= 4:
                                liability_name = row.cells[0].text.strip()
                                liability_amount = row.cells[1].text.strip()
                                asset_name = row.cells[2].text.strip()
                                asset_amount = row.cells[3].text.strip()
                                
                                # Map liability
                                if liability_name and liability_amount:
                                    bs_field = self._map_balance_sheet_field(liability_name.lower(), company)
                                    if bs_field:
                                        balance_sheet_data[bs_field] = self._parse_decimal(liability_amount)
                                        logger.info(f"DEBUG: BS Liability '{liability_name}' -> {bs_field} = {liability_amount}")
                                
                                # Map asset
                                if asset_name and asset_amount:
                                    bs_field = self._map_balance_sheet_field(asset_name.lower(), company)
                                    if bs_field:
                                        balance_sheet_data[bs_field] = self._parse_decimal(asset_amount)
                                        logger.info(f"DEBUG: BS Asset '{asset_name}' -> {bs_field} = {asset_amount}")
                    
                    elif 'expenses' in ' '.join(header_row) and 'income' in ' '.join(header_row):
                        # Profit & Loss table
                        logger.info(f"DEBUG: Parsing Profit & Loss table")
                        for row_idx in range(1, len(table.rows)):
                            row = table.rows[row_idx]
                            if len(row.cells) >= 4:
                                expense_name = row.cells[0].text.strip()
                                expense_amount = row.cells[1].text.strip()
                                income_name = row.cells[2].text.strip()
                                income_amount = row.cells[3].text.strip()
                                
                                # Map expense
                                if expense_name and expense_amount:
                                    expense_lower = expense_name.lower()
                                    # Handle "Provisions Made" separately
                                    if 'provisions made' in expense_lower or ('provisions' in expense_lower and 'made' in expense_lower):
                                        profit_loss_data['provisions'] = self._parse_decimal(expense_amount)
                                    elif 'net profit' in expense_lower:
                                        profit_loss_data['net_profit'] = self._parse_decimal(expense_amount)
                                    else:
                                        pl_field = self._map_profit_loss_field(expense_lower, is_income=False, company=company)
                                        if pl_field:
                                            profit_loss_data[pl_field] = self._parse_decimal(expense_amount)
                                            logger.info(f"DEBUG: PL Expense '{expense_name}' -> {pl_field} = {expense_amount}")
                                
                                # Map income
                                if income_name and income_amount:
                                    income_lower = income_name.lower()
                                    pl_field = self._map_profit_loss_field(income_lower, is_income=True, company=company)
                                    if pl_field:
                                        profit_loss_data[pl_field] = self._parse_decimal(income_amount)
                                        logger.info(f"DEBUG: PL Income '{income_name}' -> {pl_field} = {income_amount}")
                                
                                # Handle Net Profit if in income column
                                if income_name and 'net profit' in income_name.lower() and income_amount:
                                    profit_loss_data['net_profit'] = self._parse_decimal(income_amount)
                
                # Table 2/3: Trading Account or Operational Metrics - 2 columns (Item/Metric, Amount/Value)
                elif num_cols == 2:
                    header_row = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    logger.info(f"DEBUG: Table {table_idx + 1} header: {header_row}")
                    
                    # Check header to identify table type
                    if 'item' in ' '.join(header_row) or 'trading' in ' '.join(header_row):
                        # Trading Account table
                        logger.info(f"DEBUG: Parsing Trading Account table")
                        for row_idx in range(1, len(table.rows)):
                            row = table.rows[row_idx]
                            if len(row.cells) >= 2:
                                item_name = row.cells[0].text.strip()
                                item_amount = row.cells[1].text.strip()
                                
                                if item_name and item_amount:
                                    ta_field = self._map_trading_account_field(item_name.lower(), company)
                                    if ta_field:
                                        trading_account_data[ta_field] = self._parse_decimal(item_amount)
                                        logger.info(f"DEBUG: TA '{item_name}' -> {ta_field} = {item_amount}")
                    
                    elif 'metric' in ' '.join(header_row) or 'staff' in ' '.join(header_row):
                        # Operational Metrics table
                        logger.info(f"DEBUG: Parsing Operational Metrics table")
                        for row_idx in range(1, len(table.rows)):
                            row = table.rows[row_idx]
                            if len(row.cells) >= 2:
                                metric_name = row.cells[0].text.strip().lower()
                                metric_value = row.cells[1].text.strip()
                                
                                if 'staff' in metric_name and 'count' in metric_name and metric_value:
                                    try:
                                        operational_metrics_data['staff_count'] = int(float(metric_value.replace(',', '')))
                                        logger.info(f"DEBUG: OM Staff Count = {operational_metrics_data['staff_count']}")
                                    except:
                                        operational_metrics_data['staff_count'] = 1
            
            # Apply defaults
            balance_sheet_data = self._default_balance_sheet(balance_sheet_data)
            profit_loss_data = self._default_profit_loss(profit_loss_data)
            trading_account_data = self._default_trading_account(trading_account_data)
            if 'staff_count' not in operational_metrics_data:
                operational_metrics_data['staff_count'] = 1
            
            logger.info(f"DEBUG: Parsed .docx - BS: {len(balance_sheet_data)}, PL: {len(profit_loss_data)}, TA: {len(trading_account_data)}, OM: {operational_metrics_data}")
            
            return balance_sheet_data, profit_loss_data, trading_account_data, operational_metrics_data
            
        except Exception as e:
            logger.exception(f"Error parsing .docx file: {str(e)}")
            # Return defaults instead of empty dicts to avoid null constraint violations
            return self._default_balance_sheet({}), self._default_profit_loss({}), self._default_trading_account({}), {'staff_count': 1}

    def _parse_pdf_table(self, uploaded_file, company):
        """Parse PDF file with table format: Field | Value columns. Returns dicts for balance_sheet, profit_loss, trading_account, operational_metrics."""
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            logger.info(f"=== STARTING PDF PARSING ===")
            logger.info(f"DEBUG: File name: {uploaded_file.name if hasattr(uploaded_file, 'name') else 'Unknown'}")
            
            # Use pdfplumber to extract tables
            field_value_map = {}
            
            with pdfplumber.open(uploaded_file) as pdf:
                logger.info(f"DEBUG: PDF has {len(pdf.pages)} pages")
                
                # Try to find table on each page
                for page_num, page in enumerate(pdf.pages):
                    logger.info(f"DEBUG: Processing page {page_num + 1}")
                    # Extract tables from the page
                    tables = page.extract_tables()
                    logger.info(f"DEBUG: Found {len(tables)} tables on page {page_num + 1}")
                    
                    if tables:
                        # Use the first table found
                        table = tables[0]
                        logger.info(f"DEBUG: Using first table on page {page_num + 1} with {len(table)} rows")
                        logger.info(f"DEBUG: First few rows of table: {table[:3] if len(table) >= 3 else table}")
                        
                        # Parse table: Field | Value format
                        # Skip header row if it exists (check if first row contains "Field" or "Value")
                        start_row = 0
                        if len(table) > 0:
                            first_row_text = ' '.join([str(cell) if cell else '' for cell in table[0]]).lower()
                            logger.info(f"DEBUG: First row text: '{first_row_text}'")
                            if 'field' in first_row_text or 'value' in first_row_text:
                                start_row = 1
                                logger.info(f"DEBUG: Detected header row, starting from row {start_row}")
                        
                        # Detect table format: Expenses/Income side-by-side (4 columns) or simple Field/Value (2 columns)
                        num_cols = len(table[0]) if table else 0
                        logger.info(f"DEBUG: Table has {num_cols} columns")
                        
                        # Check if it's Expenses | Amount | Income | Amount format (4 columns)
                        is_expenses_income_format = False
                        if num_cols >= 4:
                            first_row_lower = ' '.join([str(cell).lower() if cell else '' for cell in table[0]])
                            if 'expenses' in first_row_lower and 'income' in first_row_lower:
                                is_expenses_income_format = True
                                logger.info(f"DEBUG: Detected Expenses/Income side-by-side format")
                        
                        logger.info(f"DEBUG: Parsing rows {start_row} to {len(table)-1}")
                        extracted_count = 0
                        skipped_count = 0
                        
                        if is_expenses_income_format:
                            # Format: Expenses | Amount | Income | Amount (4 columns)
                            # Or: Expenses section, then Income section (stacked)
                            logger.info(f"DEBUG: Parsing Expenses/Income format")
                            for row_idx in range(start_row, len(table)):
                                row = table[row_idx]
                                logger.debug(f"DEBUG: Row {row_idx}: {row}")
                                
                                # Check if row has Expenses data (columns 0-1)
                                if len(row) >= 2:
                                    expense_name = str(row[0]).strip() if row[0] else ''
                                    expense_value = str(row[1]).strip() if row[1] else ''
                                    
                                    if expense_name and expense_value:
                                        expense_name = expense_name.replace('\n', ' ').strip().lower()
                                        expense_value_clean = expense_value.replace(',', '').replace(' ', '').replace('-', '')
                                        if expense_value_clean.replace('.', '', 1).isdigit() and expense_name not in ['expenses', 'amount', '']:
                                            field_value_map[expense_name] = expense_value
                                            extracted_count += 1
                                            logger.info(f"DEBUG: ✓ Extracted Expense [{extracted_count}]: '{expense_name}' = '{expense_value}'")
                                
                                # Check if row has Income data (columns 2-3)
                                if len(row) >= 4:
                                    income_name = str(row[2]).strip() if row[2] else ''
                                    income_value = str(row[3]).strip() if row[3] else ''
                                    
                                    if income_name and income_value:
                                        income_name = income_name.replace('\n', ' ').strip().lower()
                                        income_value_clean = income_value.replace(',', '').replace(' ', '').replace('-', '')
                                        if income_value_clean.replace('.', '', 1).isdigit() and income_name not in ['income', 'amount', '']:
                                            field_value_map[income_name] = income_value
                                            extracted_count += 1
                                            logger.info(f"DEBUG: ✓ Extracted Income [{extracted_count}]: '{income_name}' = '{income_value}'")
                        else:
                            # Format: Field | Value (2 columns) - original logic
                            for row_idx in range(start_row, len(table)):
                                row = table[row_idx]
                                logger.debug(f"DEBUG: Row {row_idx}: {row}")
                                
                                if len(row) >= 2:
                                    field_name = str(row[0]).strip() if row[0] else ''
                                    value_str = str(row[1]).strip() if row[1] else ''
                                    
                                    # Clean up field name and value
                                    field_name = field_name.replace('\n', ' ').strip()
                                    value_str = value_str.replace('\n', ' ').strip()
                                    
                                    logger.debug(f"DEBUG: Row {row_idx} - Field: '{field_name}', Value: '{value_str}'")
                                    
                                    # Skip empty rows or header rows
                                    if field_name and value_str and field_name.lower() not in ['field', 'value', 'expenses', 'income', 'amount', '']:
                                        # Validate that value looks like a number (allow commas and decimals)
                                        value_clean = value_str.replace(',', '').replace(' ', '').replace('-', '')
                                        # Check if it's a valid number (integer or decimal)
                                        if value_clean.replace('.', '', 1).isdigit():
                                            field_value_map[field_name.lower()] = value_str
                                            extracted_count += 1
                                            logger.info(f"DEBUG: ✓ Extracted [{extracted_count}]: '{field_name}' = '{value_str}'")
                                        else:
                                            skipped_count += 1
                                            logger.warning(f"DEBUG: ✗ Skipped (not a number): '{field_name}' = '{value_str}' (cleaned: '{value_clean}')")
                                    else:
                                        skipped_count += 1
                                        logger.debug(f"DEBUG: ✗ Skipped (empty/header): Field='{field_name}', Value='{value_str}'")
                                else:
                                    skipped_count += 1
                                    logger.warning(f"DEBUG: ✗ Skipped (insufficient columns): Row has {len(row)} columns, need at least 2")
                        
                        logger.info(f"DEBUG: Extraction summary - Extracted: {extracted_count}, Skipped: {skipped_count}")
                        
                        # If we found a table, break (use first table found)
                        if field_value_map:
                            break
                
                # If no tables found, try text extraction and parse manually
                if not field_value_map:
                    logger.warning("DEBUG: ⚠️ No tables found in PDF, trying text extraction fallback")
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        logger.info(f"DEBUG: Page {page_num + 1} text length: {len(text) if text else 0} characters")
                        if text:
                            logger.info(f"DEBUG: First 500 chars of page {page_num + 1}: {text[:500]}")
                            # Try to parse Field: Value format from text
                            lines = text.split('\n')
                            logger.info(f"DEBUG: Found {len(lines)} lines in page {page_num + 1}")
                            extracted_from_text = 0
                            for line_idx, line in enumerate(lines):
                                line = line.strip()
                                if not line:
                                    continue
                                
                                # Try different separators: colon, tab, or multiple spaces
                                if ':' in line:
                                    parts = line.split(':', 1)
                                elif '\t' in line:
                                    parts = line.split('\t', 1)
                                elif '  ' in line:  # Multiple spaces (common in PDF tables)
                                    parts = line.split('  ', 1)
                                    # Clean up parts
                                    parts = [p.strip() for p in parts if p.strip()]
                                    if len(parts) < 2:
                                        continue
                                else:
                                    continue
                                
                                if len(parts) >= 2:
                                    field_name = parts[0].strip()
                                    value_str = parts[1].strip()
                                    # Remove any remaining spaces/formatting from value
                                    value_str = ' '.join(value_str.split())
                                    
                                    if field_name and value_str and field_name.lower() not in ['field', 'value']:
                                        # Check if value looks like a number
                                        value_clean = value_str.replace(',', '').replace(' ', '')
                                        if value_clean.replace('.', '').replace('-', '').isdigit():
                                            field_value_map[field_name.lower()] = value_str
                                            extracted_from_text += 1
                                            logger.info(f"DEBUG: ✓ Extracted from text [line {line_idx}]: '{field_name}' = '{value_str}'")
                            
                            logger.info(f"DEBUG: Extracted {extracted_from_text} fields from page {page_num + 1} text")
                            if extracted_from_text > 0:
                                break  # Stop after first page with data
            
            if not field_value_map:
                logger.warning("No data found in PDF file")
                return {}, {}, {}, {}
            
            logger.info(f"=== PDF EXTRACTION COMPLETE ===")
            logger.info(f"DEBUG: Total fields extracted: {len(field_value_map)}")
            logger.info(f"DEBUG: All extracted fields: {field_value_map}")
            
            if not field_value_map:
                logger.error("DEBUG: ⚠️ NO FIELDS EXTRACTED FROM PDF!")
                return {}, {}, {}, {}
            
            # Separate into different statement types (same logic as docx)
            balance_sheet_data = {}
            profit_loss_data = {}
            trading_account_data = {}
            operational_metrics_data = {}
            
            logger.info(f"=== STARTING FIELD MAPPING ===")
            # Map fields to appropriate categories (same logic as docx parsing)
            for field_lower, value_str in field_value_map.items():
                logger.info(f"DEBUG: Processing field '{field_lower}' = '{value_str}'")
                # Handle "Provisions Made" first - this is P&L provisions, different from Balance Sheet provisions
                if 'provisions made' in field_lower or ('provisions' in field_lower and 'made' in field_lower):
                    profit_loss_data['provisions'] = self._parse_decimal(value_str)
                    continue
                
                # Balance Sheet fields (check before P&L to avoid conflicts)
                bs_field = self._map_balance_sheet_field(field_lower, company)
                if bs_field:
                    balance_sheet_data[bs_field] = self._parse_decimal(value_str)
                    logger.info(f"DEBUG: ✓ Mapped '{field_lower}' -> BalanceSheet.{bs_field} = {value_str}")
                    continue
                
                # Trading Account fields (check before P&L since Gross Profit is in Trading Account)
                ta_field = self._map_trading_account_field(field_lower, company)
                if ta_field:
                    trading_account_data[ta_field] = self._parse_decimal(value_str)
                    logger.info(f"DEBUG: ✓ Mapped '{field_lower}' -> TradingAccount.{ta_field} = {value_str}")
                    continue
                else:
                    logger.debug(f"DEBUG: ✗ '{field_lower}' not mapped to TradingAccount")
                
                # Handle Gross Profit - it's calculated from Trading Account, skip if provided
                if 'gross profit' in field_lower:
                    # Gross profit is calculated, skip as it's a computed field
                    continue
                
                # Profit & Loss fields
                # Determine if income or expense based on field name patterns
                income_keywords = ['interest on loans', 'loans inter', 'interest on bank', 'miscellaneous', 'miscellane', 'return on investment', 'return on']
                expense_keywords = ['interest on deposits', 'interest on borrowings', 'borrowing', 'establishment', 'establishm', 'provisions']
                
                # Check for income patterns first
                is_income = any(x in field_lower for x in income_keywords)
                is_expense = any(x in field_lower for x in expense_keywords)
                
                # Special handling for "Miscellaneous" / "Miscellane" - always income
                if 'miscellaneous' in field_lower or 'miscellane' in field_lower:
                    is_income = True
                    is_expense = False
                    logger.info(f"DEBUG: Detected Miscellaneous/Miscellane as income")
                
                # Special handling for "Interest o" / "interest_o" - could be income or expense
                if ('interest o' in field_lower or 'interest_o' in field_lower) and not is_expense:
                    # If it's in the income section or matches income patterns, treat as income
                    if 'loan' in field_lower or 'bank' in field_lower:
                        is_income = True
                        logger.info(f"DEBUG: Detected 'Interest o' as income (loans/bank context)")
                    elif 'deposit' in field_lower or 'borrowing' in field_lower:
                        is_expense = True
                        logger.info(f"DEBUG: Detected 'Interest o' as expense (deposits/borrowings context)")
                
                # Handle Net Profit separately (it's in P&L but not income/expense)
                if 'net profit' in field_lower:
                    profit_loss_data['net_profit'] = self._parse_decimal(value_str)
                    logger.info(f"DEBUG: ✓ Mapped '{field_lower}' -> ProfitLoss.net_profit = {value_str}")
                    continue
                
                pl_field = self._map_profit_loss_field(field_lower, is_income=is_income, company=company)
                if pl_field:
                    profit_loss_data[pl_field] = self._parse_decimal(value_str)
                    logger.info(f"DEBUG: ✓ Mapped '{field_lower}' -> ProfitLoss.{pl_field} = {value_str} (is_income={is_income})")
                    continue
                else:
                    logger.warning(f"DEBUG: ✗ '{field_lower}' not mapped to ProfitLoss (is_income={is_income}, is_expense={is_expense})")
                
                # Operational Metrics
                if 'staff' in field_lower and 'count' in field_lower:
                    try:
                        operational_metrics_data['staff_count'] = int(float(value_str.replace(',', '')))
                        logger.info(f"DEBUG: ✓ Mapped '{field_lower}' -> OperationalMetrics.staff_count = {value_str}")
                    except:
                        operational_metrics_data['staff_count'] = 1
                        logger.warning(f"DEBUG: ✗ Failed to parse staff_count, using default: 1")
                    continue
                
                # If we get here, the field wasn't mapped to any category
                logger.warning(f"DEBUG: ⚠️ UNMAPPED FIELD: '{field_lower}' = '{value_str}'")
            
            logger.info(f"=== FIELD MAPPING COMPLETE ===")
            logger.info(f"DEBUG: Balance Sheet fields mapped: {len(balance_sheet_data)} - {list(balance_sheet_data.keys())}")
            logger.info(f"DEBUG: Profit & Loss fields mapped: {len(profit_loss_data)} - {list(profit_loss_data.keys())}")
            logger.info(f"DEBUG: Trading Account fields mapped: {len(trading_account_data)} - {list(trading_account_data.keys())}")
            logger.info(f"DEBUG: Operational Metrics: {operational_metrics_data}")
            
            # Apply defaults BEFORE returning
            balance_sheet_data = self._default_balance_sheet(balance_sheet_data)
            profit_loss_data = self._default_profit_loss(profit_loss_data)
            trading_account_data = self._default_trading_account(trading_account_data)
            if 'staff_count' not in operational_metrics_data:
                operational_metrics_data['staff_count'] = 1
            
            logger.info(f"=== FINAL RESULTS AFTER DEFAULTS ===")
            logger.info(f"DEBUG: Balance Sheet: {len(balance_sheet_data)} fields")
            logger.info(f"DEBUG: Profit & Loss: {len(profit_loss_data)} fields")
            logger.info(f"DEBUG: Profit & Loss values: interest_on_loans={profit_loss_data.get('interest_on_loans')}, interest_on_bank_ac={profit_loss_data.get('interest_on_bank_ac')}, miscellaneous_income={profit_loss_data.get('miscellaneous_income')}, return_on_investment={profit_loss_data.get('return_on_investment')}")
            logger.info(f"DEBUG: Trading Account: {len(trading_account_data)} fields")
            logger.info(f"DEBUG: Trading Account values: opening_stock={trading_account_data.get('opening_stock')}, purchases={trading_account_data.get('purchases')}, sales={trading_account_data.get('sales')}, closing_stock={trading_account_data.get('closing_stock')}, trade_charges={trading_account_data.get('trade_charges')}")
            logger.info(f"DEBUG: Operational Metrics: {operational_metrics_data}")
            
            # Final safety check - ensure miscellaneous_income is never null
            if 'miscellaneous_income' not in profit_loss_data or profit_loss_data.get('miscellaneous_income') is None:
                profit_loss_data['miscellaneous_income'] = Decimal('0')
                logger.warning(f"DEBUG: ⚠️ miscellaneous_income was null, set to 0")
            
            return balance_sheet_data, profit_loss_data, trading_account_data, operational_metrics_data
            
        except Exception as e:
            logger.exception(f"Error parsing PDF file: {str(e)}")
            # Return defaults instead of empty dicts to avoid null constraint violations
            return self._default_balance_sheet({}), self._default_profit_loss({}), self._default_trading_account({}), {'staff_count': 1}

    def _create_period_from_document(self, request, company, uploaded_file, period_info, file_type):
        """Create a financial period for .docx or .pdf upload; store file and create/update records."""
        period_label = request.data.get('period_label') or period_info.get('label') or f"FY-{datetime.now().year}-{datetime.now().year + 1}"
        start_date = request.data.get('start_date') or period_info.get('start_date') or f"{datetime.now().year}-04-01"
        end_date = request.data.get('end_date') or period_info.get('end_date') or f"{datetime.now().year + 1}-03-31"
        period_type = request.data.get('period_type') or period_info.get('period_type') or 'YEARLY'

        with transaction.atomic():
            period, created = FinancialPeriod.objects.get_or_create(
                company=company,
                label=period_label,
                defaults={
                    'period_type': period_type,
                    'start_date': start_date,
                    'end_date': end_date,
                    'is_finalized': False,
                }
            )
            
            # Parse .docx file if it's a docx upload
            if file_type == 'docx':
                # Parse the .docx table and extract data BEFORE saving (to avoid file pointer issues)
                balance_sheet_data, profit_loss_data, trading_account_data, operational_metrics_data = self._parse_docx_table(uploaded_file, company)
                
                # Ensure defaults are applied before saving (safety check)
                trading_account_data = self._default_trading_account(trading_account_data)
                balance_sheet_data = self._default_balance_sheet(balance_sheet_data)
                profit_loss_data = self._default_profit_loss(profit_loss_data)
                
                logger.info(f"DEBUG: Before save (docx) - Trading Account data: {trading_account_data}")
                
                # Reset file pointer after parsing (Django needs it for saving)
                uploaded_file.seek(0)
                period.uploaded_file = uploaded_file
                period.file_type = 'docx'
                period.save()
                
                # Ensure staff_count is always set before saving (safety check)
                if 'staff_count' not in operational_metrics_data or operational_metrics_data.get('staff_count') is None:
                    operational_metrics_data['staff_count'] = 1
                
                # Create/update records with parsed data
                TradingAccount.objects.update_or_create(period=period, defaults=trading_account_data)
                ProfitAndLoss.objects.update_or_create(period=period, defaults=profit_loss_data)
                BalanceSheet.objects.update_or_create(period=period, defaults=balance_sheet_data)
                OperationalMetrics.objects.update_or_create(period=period, defaults=operational_metrics_data)
                
                # Calculate ratios
                from app.services.ratio_calculator import RatioCalculator
                calculator = RatioCalculator(period)
                all_ratios = calculator.calculate_all_ratios()
                traffic_light_statuses = calculator.get_traffic_light_statuses()
                
                RatioResult.objects.update_or_create(
                    period=period,
                    defaults={
                        'working_fund': Decimal(str(all_ratios['working_fund'])),
                        'stock_turnover': Decimal(str(all_ratios.get('stock_turnover', 0))),
                        'gross_profit_ratio': Decimal(str(all_ratios.get('gross_profit_ratio', 0))),
                        'net_profit_ratio': Decimal(str(all_ratios.get('net_profit_ratio', 0))),
                        'own_fund_to_wf': Decimal(str(all_ratios.get('own_fund_to_wf', 0))),
                        'deposits_to_wf': Decimal(str(all_ratios.get('deposits_to_wf', 0))),
                        'borrowings_to_wf': Decimal(str(all_ratios.get('borrowings_to_wf', 0))),
                        'loans_to_wf': Decimal(str(all_ratios.get('loans_to_wf', 0))),
                        'investments_to_wf': Decimal(str(all_ratios.get('investments_to_wf', 0))),
                        'cost_of_deposits': Decimal(str(all_ratios.get('cost_of_deposits', 0))),
                        'yield_on_loans': Decimal(str(all_ratios.get('yield_on_loans', 0))),
                        'yield_on_investments': Decimal(str(all_ratios.get('yield_on_investments', 0))),
                        'credit_deposit_ratio': Decimal(str(all_ratios.get('credit_deposit_ratio', 0))),
                        'avg_cost_of_wf': Decimal(str(all_ratios.get('avg_cost_of_wf', 0))),
                        'avg_yield_on_wf': Decimal(str(all_ratios.get('avg_yield_on_wf', 0))),
                        'gross_fin_margin': Decimal(str(all_ratios.get('gross_fin_margin', 0))),
                        'operating_cost_to_wf': Decimal(str(all_ratios.get('operating_cost_to_wf', 0))),
                        'net_fin_margin': Decimal(str(all_ratios.get('net_fin_margin', 0))),
                        'risk_cost_to_wf': Decimal(str(all_ratios.get('risk_cost_to_wf', 0))),
                        'net_margin': Decimal(str(all_ratios.get('net_margin', 0))),
                        'all_ratios': all_ratios,
                        'traffic_light_status': traffic_light_statuses
                    }
                )
            else:
                # PDF: parse file and extract data
                # Parse the PDF table and extract data BEFORE saving (to avoid file pointer issues)
                balance_sheet_data, profit_loss_data, trading_account_data, operational_metrics_data = self._parse_pdf_table(uploaded_file, company)
                
                # Ensure defaults are applied before saving (safety check)
                trading_account_data = self._default_trading_account(trading_account_data)
                balance_sheet_data = self._default_balance_sheet(balance_sheet_data)
                profit_loss_data = self._default_profit_loss(profit_loss_data)
                
                logger.info(f"DEBUG: Before save - Trading Account data: {trading_account_data}")
                
                # Reset file pointer after parsing (Django needs it for saving)
                uploaded_file.seek(0)
                period.uploaded_file = uploaded_file
                period.file_type = 'pdf'
                period.save()
                
                # Ensure staff_count is always set before saving (safety check)
                if 'staff_count' not in operational_metrics_data or operational_metrics_data.get('staff_count') is None:
                    operational_metrics_data['staff_count'] = 1
                
                # Create/update records with parsed data
                TradingAccount.objects.update_or_create(period=period, defaults=trading_account_data)
                ProfitAndLoss.objects.update_or_create(period=period, defaults=profit_loss_data)
                BalanceSheet.objects.update_or_create(period=period, defaults=balance_sheet_data)
                OperationalMetrics.objects.update_or_create(period=period, defaults=operational_metrics_data)
                
                # Calculate ratios
                from app.services.ratio_calculator import RatioCalculator
                calculator = RatioCalculator(period)
                all_ratios = calculator.calculate_all_ratios()
                traffic_light_statuses = calculator.get_traffic_light_statuses()
                
                RatioResult.objects.update_or_create(
                    period=period,
                    defaults={
                        'working_fund': Decimal(str(all_ratios['working_fund'])),
                        'stock_turnover': Decimal(str(all_ratios.get('stock_turnover', 0))),
                        'gross_profit_ratio': Decimal(str(all_ratios.get('gross_profit_ratio', 0))),
                        'net_profit_ratio': Decimal(str(all_ratios.get('net_profit_ratio', 0))),
                        'own_fund_to_wf': Decimal(str(all_ratios.get('own_fund_to_wf', 0))),
                        'deposits_to_wf': Decimal(str(all_ratios.get('deposits_to_wf', 0))),
                        'borrowings_to_wf': Decimal(str(all_ratios.get('borrowings_to_wf', 0))),
                        'loans_to_wf': Decimal(str(all_ratios.get('loans_to_wf', 0))),
                        'investments_to_wf': Decimal(str(all_ratios.get('investments_to_wf', 0))),
                        'cost_of_deposits': Decimal(str(all_ratios.get('cost_of_deposits', 0))),
                        'yield_on_loans': Decimal(str(all_ratios.get('yield_on_loans', 0))),
                        'yield_on_investments': Decimal(str(all_ratios.get('yield_on_investments', 0))),
                        'credit_deposit_ratio': Decimal(str(all_ratios.get('credit_deposit_ratio', 0))),
                        'avg_cost_of_wf': Decimal(str(all_ratios.get('avg_cost_of_wf', 0))),
                        'avg_yield_on_wf': Decimal(str(all_ratios.get('avg_yield_on_wf', 0))),
                        'gross_fin_margin': Decimal(str(all_ratios.get('gross_fin_margin', 0))),
                        'operating_cost_to_wf': Decimal(str(all_ratios.get('operating_cost_to_wf', 0))),
                        'net_fin_margin': Decimal(str(all_ratios.get('net_fin_margin', 0))),
                        'risk_cost_to_wf': Decimal(str(all_ratios.get('risk_cost_to_wf', 0))),
                        'net_margin': Decimal(str(all_ratios.get('net_margin', 0))),
                        'all_ratios': all_ratios,
                        'traffic_light_status': traffic_light_statuses
                    }
                )
        return period

    def _extract_period_from_filename(self, filename):
        """
        Extract period information from filename.
        India FY (Apr-Mar) formats: Apr_2024, Q1_FY_2024_25, H1_FY_2024_25, FY_2024_25
        Legacy: April_2025, April-2025, etc.
        """
        import re
        from datetime import datetime
        from app.services.period_labels import parse_period_label

        period_info = {}
        name_without_ext = filename.rsplit('.', 1)[0] if '.' in filename else filename

        # 1. Try India FY format first (Apr_2024, Q1_FY_2024_25, H1_FY_2024_25, FY_2024_25)
        info = parse_period_label(name_without_ext)
        if info:
            return info

        # 2. Fallback: full month names (April_2025, April-2025, etc.)
        month_names = {
            'january': 1, 'february': 2, 'march': 3, 'april': 4,
            'may': 5, 'june': 6, 'july': 7, 'august': 8,
            'september': 9, 'october': 10, 'november': 11, 'december': 12
        }
        month_patterns = [
            r'([A-Za-z]+)[_\-\s]+(\d{4})',  # April_2025, April-2025
            r'(\d{1,2})[_\-\s]+([A-Za-z]+)[_\-\s]+(\d{4})',  # 01_April_2025
            r'([A-Za-z]+)[_\-\s]+(\d{4})[_\-\s]+([A-Za-z]+)',  # April_2025_March
        ]
        for pattern in month_patterns:
            match = re.search(pattern, name_without_ext, re.IGNORECASE)
            if match:
                groups = match.groups()
                month_name = None
                month_num = None
                year = None
                for group in groups:
                    group_lower = str(group).lower()
                    if group_lower in month_names:
                        month_name = group_lower
                        month_num = month_names[month_name]
                    elif str(group).isdigit() and len(str(group)) == 4:
                        year = int(group)
                if month_name and year and month_num:
                    period_info['label'] = f"{month_name.capitalize()}_{year}"
                    start_date = datetime(year, month_num, 1)
                    if month_num == 2:
                        end_day = 29 if (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0) else 28
                    elif month_num in [4, 6, 9, 11]:
                        end_day = 30
                    else:
                        end_day = 31
                    end_date = datetime(year, month_num, end_day)
                    period_info['start_date'] = start_date.strftime('%Y-%m-%d')
                    period_info['end_date'] = end_date.strftime('%Y-%m-%d')
                    period_info['period_type'] = 'MONTHLY'
                    break
        return period_info


class RatioBenchmarksView(APIView):
    """GET: return current ratio benchmarks (DB merged with defaults). PUT: update stored benchmarks."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        from app.services.benchmark_config import get_ratio_benchmarks
        from app.config.ratio_benchmarks import DEFAULT_RATIO_BENCHMARKS
        data = get_ratio_benchmarks()
        # Include labels for frontend (same keys as defaults)
        labels = {
            "stock_turnover": "Stock Turnover (times/year)",
            "gross_profit_ratio_min": "Gross Profit Ratio Min (%)",
            "gross_profit_ratio_max": "Gross Profit Ratio Max (%)",
            "own_fund_to_wf": "Own Fund to Working Fund (%)",
            "loans_to_wf_min": "Loans to WF Min (%)",
            "loans_to_wf_max": "Loans to WF Max (%)",
            "investments_to_wf_min": "Investments to WF Min (%)",
            "investments_to_wf_max": "Investments to WF Max (%)",
            "avg_cost_of_wf": "Avg Cost of WF (%)",
            "avg_yield_on_wf": "Avg Yield on WF (%)",
            "gross_financial_margin": "Gross Financial Margin (%)",
            "operating_cost_to_wf_min": "Operating Cost to WF Min (%)",
            "operating_cost_to_wf_max": "Operating Cost to WF Max (%)",
            "net_financial_margin": "Net Financial Margin (%)",
            "risk_cost_to_wf_max": "Risk Cost to WF Max (%)",
            "net_margin": "Net Margin (%)",
            "credit_deposit_ratio_min": "Credit Deposit Ratio Min (%)",
        }
        return Response({
            "benchmarks": data,
            "labels": labels,
            "keys_order": list(DEFAULT_RATIO_BENCHMARKS.keys()),
        })

    def put(self, request):
        from app.services.benchmark_config import set_ratio_benchmarks
        try:
            data = request.data.get("benchmarks") if isinstance(request.data, dict) else request.data
            if not isinstance(data, dict):
                return Response({
                    "status": "failed",
                    "message": "benchmarks must be an object",
                }, status=status.HTTP_400_BAD_REQUEST)
            set_ratio_benchmarks(data)
            return Response({"status": "success", "message": "Benchmarks updated."})
        except ValueError as e:
            return Response({"status": "failed", "message": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({
                "status": "failed",
                "message": str(e),
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class BulkImportCompaniesView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        """
        Bulk import companies
        POST /api/companies/bulk_import/
        """
        try:
            companies_data = request.data.get('companies', [])
            if not companies_data:
                return Response({
                    "status": "failed",
                    "response_code": status.HTTP_400_BAD_REQUEST,
                    "message": "No companies provided"
                })

            success_count = 0
            failed_count = 0
            errors = []

            for company_data in companies_data:
                try:
                    name = company_data.get('name')
                    registration_no = company_data.get('registration_no')

                    if not name or not registration_no:
                        failed_count += 1
                        errors.append(f"Missing name or registration number for data: {company_data}")
                        continue

                    # Check if company with registration number already exists
                    if Company.objects.filter(registration_no=registration_no).exists():
                        failed_count += 1
                        errors.append(f"Company with registration no {registration_no} already exists")
                        continue

                    # Create company
                    Company.objects.create(
                        name=name,
                        registration_no=registration_no
                    )
                    success_count += 1

                except Exception as e:
                    failed_count += 1
                    errors.append(f"Error creating company {company_data.get('name', 'Unknown')}: {str(e)}")

            return Response({
                "status": "success",
                "response_code": status.HTTP_200_OK,
                "success": success_count,
                "failed": failed_count,
                "errors": errors,
                "message": f"Successfully imported {success_count} companies. {failed_count} failed."
            })

        except Exception as e:
            logger.exception(f"Error in BulkImportCompaniesView: {str(e)}")
            return Response({
                "status": "failed",
                "response_code": status.HTTP_500_INTERNAL_SERVER_ERROR,
                "message": str(e)
            })


class DownloadExcelTemplateView(APIView):
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        """Download Excel template with 4 sheets"""
        try:
            wb = Workbook()
            
            # Remove default sheet
            wb.remove(wb.active)
            
            # Sheet 1: Balance Sheet
            ws_bs = wb.create_sheet("Balance Sheet")
            ws_bs.append(["Liabilities", "Amount", "Assets", "Amount"])
            ws_bs.append(["Share Capital", 5281006, "Cash in Hand", 484706199])
            ws_bs.append(["Deposits", 484706199, "Cash at Bank", 90000000])
            ws_bs.append(["Borrowings", 7001911, "Investments", 13328928])
            ws_bs.append(["Reserves (Statutory & Free)", 10569840, "Loans & Advances", 437223261])
            ws_bs.append(["Provisions", 53117811, "Fixed Assets", 55501843])
            ws_bs.append(["Other Liabilities", 46444029, "Other Assets", 5678014])
            ws_bs.append(["Undistributed Profit", 10866453, "Stock in Trade", 40000])
            
            # Sheet 2: Profit & Loss
            ws_pl = wb.create_sheet("Profit & Loss")
            ws_pl.append(["Expenses", "Amount", "Income", "Amount"])
            ws_pl.append(["Interest on Deposits", 26698057, "Interest on Loans", 42488657])
            ws_pl.append(["Interest on Borrowings", 770021, "Interest on Bank A/c", 6300000])
            ws_pl.append(["Establishment & Contingencies", 13476132, "Return on Investment", 1066314])
            ws_pl.append(["Provisions Made", 4533930, "Miscellaneous Income", 3485633])
            ws_pl.append(["Net Profit", 7863516, "", ""])
            
            # Sheet 3: Trading Account
            ws_ta = wb.create_sheet("Trading Account")
            ws_ta.append(["Item", "Amount"])
            ws_ta.append(["Opening Stock", 25080])
            ws_ta.append(["Purchases", 572444])
            ws_ta.append(["Trade Charges", 8176])
            ws_ta.append(["Sales", 552264])
            ws_ta.append(["Closing Stock", 40000])
            
            # Sheet 4: Operational Metrics
            ws_om = wb.create_sheet("Operational Metrics")
            ws_om.append(["Metric", "Value"])
            ws_om.append(["Staff Count", 24])
            
            # Save to BytesIO
            output = BytesIO()
            wb.save(output)
            output.seek(0)
            
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename="Financial_Data_Template.xlsx"'
            return response
            
        except Exception as e:
            logger.error(f"Error generating Excel template: {e}")
            return Response({
                "status": "failed",
                "message": str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class DownloadWordTemplateView(APIView):
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        """Download Word template with 4 tables"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Financial Data Template', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 1. Balance Sheet
            doc.add_heading('1. Balance Sheet', level=1)
            table_bs = doc.add_table(rows=8, cols=4)
            table_bs.style = 'Light Grid Accent 1'
            
            # Header
            table_bs.cell(0, 0).text = 'Liabilities'
            table_bs.cell(0, 1).text = 'Amount'
            table_bs.cell(0, 2).text = 'Assets'
            table_bs.cell(0, 3).text = 'Amount'
            
            # Data
            bs_data = [
                ['Share Capital', '5281006', 'Cash in Hand', '484706199'],
                ['Deposits', '484706199', 'Cash at Bank', '90000000'],
                ['Borrowings', '7001911', 'Investments', '13328928'],
                ['Reserves (Statutory & Free)', '10569840', 'Loans & Advances', '437223261'],
                ['Provisions', '53117811', 'Fixed Assets', '55501843'],
                ['Other Liabilities', '46444029', 'Other Assets', '5678014'],
                ['Undistributed Profit', '10866453', 'Stock in Trade', '40000'],
            ]
            
            for i, row_data in enumerate(bs_data, start=1):
                for j, cell_text in enumerate(row_data):
                    table_bs.cell(i, j).text = cell_text
            
            doc.add_paragraph()
            
            # 2. Profit and Loss
            doc.add_heading('2. Profit and Loss', level=1)
            table_pl = doc.add_table(rows=6, cols=4)
            table_pl.style = 'Light Grid Accent 1'
            
            table_pl.cell(0, 0).text = 'Expenses'
            table_pl.cell(0, 1).text = 'Amount'
            table_pl.cell(0, 2).text = 'Income'
            table_pl.cell(0, 3).text = 'Amount'
            
            pl_data = [
                ['Interest on Deposits', '26698057', 'Interest on Loans', '42488657'],
                ['Interest on Borrowings', '770021', 'Interest on Bank A/c', '6300000'],
                ['Establishment & Contingencies', '13476132', 'Return on Investment', '1066314'],
                ['Provisions Made', '4533930', 'Miscellaneous Income', '3485633'],
                ['Net Profit', '7863516', '', ''],
            ]
            
            for i, row_data in enumerate(pl_data, start=1):
                for j, cell_text in enumerate(row_data):
                    table_pl.cell(i, j).text = cell_text
            
            doc.add_paragraph()
            
            # 3. Trading Account
            doc.add_heading('3. Trading Account', level=1)
            table_ta = doc.add_table(rows=6, cols=2)
            table_ta.style = 'Light Grid Accent 1'
            
            table_ta.cell(0, 0).text = 'Item'
            table_ta.cell(0, 1).text = 'Amount'
            
            ta_data = [
                ['Opening Stock', '25080'],
                ['Purchases', '572444'],
                ['Trade Charges', '8176'],
                ['Sales', '552264'],
                ['Closing Stock', '40000'],
            ]
            
            for i, row_data in enumerate(ta_data, start=1):
                table_ta.cell(i, 0).text = row_data[0]
                table_ta.cell(i, 1).text = row_data[1]
            
            doc.add_paragraph()
            
            # 4. Operational Metrics
            doc.add_heading('4. Operational Metrics', level=1)
            table_om = doc.add_table(rows=2, cols=2)
            table_om.style = 'Light Grid Accent 1'
            
            table_om.cell(0, 0).text = 'Metric'
            table_om.cell(0, 1).text = 'Value'
            table_om.cell(1, 0).text = 'Staff Count'
            table_om.cell(1, 1).text = '24'
            
            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename="Financial_Data_Template.docx"'
            return response
            
        except Exception as e:
            logger.error(f"Error generating Word template: {e}")
            return Response({
                "status": "failed",
                "message": str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
